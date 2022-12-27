import os
import subprocess
import time
import paramiko
import threading
from fpdf import FPDF
from docx import Document
from Apps.app_common import *
from Apps.app_ran import *
from PySimpleGUI import PySimpleGUI as sg
import re
from getpass import getpass

#################################################
#      Interfaces gráficas (installation)       #
#################################################


def interface1():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface1.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='center')],
        [sg.Text(' ')],
        [sg.Text('Welcome, this platform aims to help your experience in using OpenAirInterface.')],
        [sg.Text('Depending on your needs, choose one of the following options:')],
        [sg.Text(' ')],
        [sg.Radio('Install software', "RADIO1", key='installsoft')],
        [sg.Radio('Express Deployment (less settings are required)', "RADIO1", key='fastdeployment')],
        [sg.Radio('Custom Deployment (more settings are required)', "RADIO1", key='advancedeployment')],
        [sg.Radio('Tests', "RADIO1", key='tests')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface2():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/install_options.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('The following options allow you to choose how the OAI software will be installed')],
        [sg.Text('If you want to install only one of the network elements choose one')],
        [sg.Text('of the 3 options ("Install Core";"Install gNB";"Install OAI UE")')],
        [sg.Text('depending on what you want to install. If you want to install more')],
        [sg.Text('than one element, choose the "Complete Installation" option.')],
        [sg.Text(' ')],
        [sg.Radio('Install Core', "Radio1",key='installcore')],
        [sg.Radio('Install gNB', "RADIO1", key='installgnb')],
        [sg.Radio('Install OAI UE', "RADIO1", key='installoaiue')],
        [sg.Radio('Complete Installation',"RADIO1", key='installcomplete')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'),sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface3():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/uetype.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('If you are going to use a COTS UE terminal choose the option "COTS UE",')],
        [sg.Text('however, if you also want to install the OAI UE terminal, choose the "OAI UE" option.')],
        [sg.Text('Choose terminal type:')],
        [sg.Radio('OAI UE', "RADIO1", key='oaiue')],
        [sg.Radio('COTS UE', "RADIO1", key='realue')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface4():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setuptype.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Choose installation type')],
        [sg.Radio('All-in-one(Core+gNB on same machine) and OAI UE', "RADIO1", key='allinoneoaiue')],
        [sg.Radio('Core and gNB in different machines and OAI UE', "RADIO1", key='coregnboaiue')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface5():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface5_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Choose installation type')],
        [sg.Radio('All-in-one(Core+gNB on same machine)', "RADIO1", key='allinonerealue')],
        [sg.Radio('Core and gNB on diferent machines', "RADIO1", key='coregnbrealue')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface6():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/core_information.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Information about the machine where the core will be installed')],
        [sg.Text(' ')],
        [sg.Text('Core data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcore',password_char='*')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface7():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/gnb_information.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Information about the machine where the gNB will be installed')],
        [sg.Text(' ')],
        [sg.Text('gNB data')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipgnb')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usergnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passgnb',password_char='*')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface8():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/ue_information.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Information about the machine where the OAI UE will be installed')],
        [sg.Text(' ')],
        [sg.Text('OAI UE data')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipoaiue')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='useroaiue')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passoaiue',password_char='*')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface9():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation1.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Information about the machines where the core, gNb and OAI UE will be installed')],
        [sg.Text('Core and gNB data')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcoregnb')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercoregnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcoregnb',password_char='*')],
        [sg.Text('OAI UE data')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipoaiue')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='useroaiue')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passoaiue',password_char='*')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface10():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation2.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Information about the machines where the core, gNb and OAI UE will be installed')],
        [sg.Text('Core data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcore',password_char='*')],
        [sg.Text('gNB data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipgnb')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usergnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passgnb',password_char='*')],
        [sg.Text('OAI UE data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipoaiue')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='useroaiue')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passoaiue',password_char='*')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 620))


def interface11():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation3.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Information about the machines where the core, gNb and OAI UE will be installed')],
        [sg.Text('Core and gNB data')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcoregnb')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercoregnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcoregnb',password_char='*')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface12():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation4.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Information about the machines where the core and gNb will be installed')],
        [sg.Text('Core data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcore',password_char='*')],
        [sg.Text('gNB data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipgnb')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usergnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passgnb',password_char='*')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface13():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    my_img = sg.Image(filename='Images/interface13_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Installing Core')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Column([[gif]], justification='center')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface14():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    my_img = sg.Image(filename='Images/interface14_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Installing gNB')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Column([[gif]], justification='center')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface15():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    my_img = sg.Image(filename='Images/interface15_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Installing OAI UE')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Column([[gif]], justification='center')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface16():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    my_img = sg.Image(filename='Images/interface16_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Installing All-in-One and OAI UE')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Column([[gif]], justification='center')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface17():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    my_img = sg.Image(filename='Images/interface17_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Installing Core and gNB and OAI UE')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Column([[gif]], justification='center')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface18():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    my_img = sg.Image(filename='Images/interface18_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Installing All-in-One')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Column([[gif]], justification='center')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface19():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    my_img = sg.Image(filename='Images/interface19_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Installing Core and gNB')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Column([[gif]], justification='center')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface20():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/installationcorecompleted.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('Core installation is complete')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Close')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface21():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/InstallationgNBcomplete.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('gNB installation is complete')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Close')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface22():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/installationoaiuecompleted.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('OAI UE Terminal installation is complete')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Close')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface23():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/installationallinoneoaiuecompleted.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('The installation of All-in-one (Core + gNB on same machine) is complete.')],
        [sg.Text('OAI UE Terminal installation is complete')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Close')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface24():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/installationcoregnboaiuecompleted.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Core installation is complete')],
        [sg.Text('gNB installation is complete')],
        [sg.Text('OAI UE Terminal installation is complete')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Close')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface25():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/installationallinonecompleted.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('The installation of All-in-one (Core + gNB on same machine) is complete.')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Close')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface26():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/installationcoregnbcompleted.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('Core installation is complete')],
        [sg.Text('gNB installation is complete')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Close')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))

#################################################
#      Interfaces gráficas (installation)       #
#################################################


def interface27():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface27_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('What type of implementation is used?')],
        [sg.Text(' ')],
        [sg.Radio('All-in-One (Core + gNB on same machine)', "RADIO1", key='allinone')],
        [sg.Radio('Core and gNB on different machines', "RADIO1", key='coregnb')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 620))


def interface28():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/uetype.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('What type of UE used?')],
        [sg.Text(' ')],
        [sg.Radio('OAI UE', "RADIO1", key='oaiue')],
        [sg.Radio('COTS UE', "RADIO1", key='realue')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface29():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation1.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Core and gNB data:', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcoregnb'), sg.Text("", size=(0, 1), key='last_ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercoregnb'), sg.Text("", size=(0, 1), key='last_usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcoregnb',password_char='*'), sg.Text("", size=(0, 1), key='last_passcore')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dircoregnb'), sg.Text("", size=(0, 1), key='last_dircore')],
        [sg.Text('OAI UE data')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipoaiue'), sg.Text("", size=(0, 1), key='last_ipoaiue')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='useroaiue'), sg.Text("", size=(0, 1), key='last_useroaiue')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passoaiue',password_char='*'), sg.Text("", size=(0, 1), key='last_passoaiue')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirue'), sg.Text("", size=(0, 1), key='last_diroaiue')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 650))


def interface30():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation_interface30.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Core and gNB data:', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcoregnb'), sg.Text("", size=(0, 1), key='last_ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercoregnb'), sg.Text("", size=(0, 1), key='last_usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcoregnb',password_char='*'), sg.Text("", size=(0, 1), key='last_passcore')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dircoregnb'), sg.Text("", size=(0, 1), key='last_dircore')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface31():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation_interface31.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Core data:', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcore'), sg.Text("", size=(0, 1), key='last_ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercore'), sg.Text("", size=(0, 1), key='last_usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcore',password_char='*'), sg.Text("", size=(0, 1), key='last_passcore')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dircore'), sg.Text("", size=(0, 1), key='last_dircore')],
        [sg.Text('gNB data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipgnb'), sg.Text("", size=(0, 1), key='last_ipgnb')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usergnb'), sg.Text("", size=(0, 1), key='last_usergnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passgnb',password_char='*'), sg.Text("", size=(0, 1), key='last_passgnb')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirgnb'), sg.Text("", size=(0, 1), key='last_dirgnb')],
        [sg.Text('OAI UE data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipoaiue'), sg.Text("", size=(0, 1), key='last_ipoaiue')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='useroaiue'), sg.Text("", size=(0, 1), key='last_useroaiue')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passoaiue',password_char='*'), sg.Text("", size=(0, 1), key='last_passoaiue')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirue'), sg.Text("", size=(0, 1), key='last_diroaiue')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 750))


def interface32():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation_interface32.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Core data:', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcore'), sg.Text("", size=(0, 1), key='last_ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercore'), sg.Text("", size=(0, 1), key='last_usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcore',password_char='*'), sg.Text("", size=(0, 1), key='last_passcore')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dircore'), sg.Text("", size=(0, 1), key='last_dircore')],
        [sg.Text('gNB data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipgnb'), sg.Text("", size=(0, 1), key='last_ipgnb')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usergnb'), sg.Text("", size=(0, 1), key='last_usergnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passgnb',password_char='*'), sg.Text("", size=(0, 1), key='last_passgnb')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirgnb'), sg.Text("", size=(0, 1), key='last_dirgnb')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface33():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/sim_card_programming.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Program the SIM Card')],
        [sg.Text('ADM: '), sg.Input('', key='adm')],
        [sg.Text('The SIM CARD will be programmed with the following data:')],
        [sg.Text('imsi = 208990000000001')],
        [sg.Text('key = fec86ba6eb707ed08905757b1bb44b8f')],
        [sg.Text('opc = C42449363BBAD02B66D16BC975D77CC1')],
        [sg.Text('spn = OpenAirInterface')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface34():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    layout = [
        [sg.Text('Preparing Setup....')],
        [sg.Text('Preparing configuration files and connecting setup')],
        [sg.Text('Setup Type: '), sg.Text("", size=(0, 1), key='setuptype')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Column([[gif]], justification='center')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface35():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface35_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Identifies the NIC of each machine:')],
        [sg.Text('NIC Core:')],
        [sg.Combo(nic_core, key='niccore')],
        [sg.Text('NIC gNB:')],
        [sg.Combo(nic_gnb, key='nicgnb')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))

###############################################
#                   Advance                   #
###############################################


def interface36():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface36_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('What type of implementation is used?')],
        [sg.Text(' ')],
        [sg.Radio('All-in-One (Core + gNB on same machine)', "RADIO1", key='allinone')],
        [sg.Radio('Core and gNB on different machines', "RADIO1", key='coregnb')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 620))


def interface37():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/uetype.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('What type of UE used?')],
        [sg.Text(' ')],
        [sg.Radio('OAI UE', "RADIO1", key='oaiue')],
        [sg.Radio('COTS UE', "RADIO1", key='realue')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface38():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation1.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Core and gNB data:', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcoregnb'), sg.Text("", size=(0, 1), key='last_ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercoregnb'), sg.Text("", size=(0, 1), key='last_usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcoregnb',password_char='*'), sg.Text("", size=(0, 1), key='last_passcore')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dircoregnb'), sg.Text("", size=(0, 1), key='last_dircore')],
        [sg.Text('OAI UE data')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipoaiue'), sg.Text("", size=(0, 1), key='last_ipoaiue')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='useroaiue'), sg.Text("", size=(0, 1), key='last_useroaiue')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passoaiue',password_char='*'), sg.Text("", size=(0, 1), key='last_passoaiue')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirue'), sg.Text("", size=(0, 1), key='last_diroaiue')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 630))


def interface39():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation_interface30.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Core and gNB data:', size=(52, 1)), sg.Text('Previous information:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcoregnb'), sg.Text("", size=(0, 1), key='last_ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercoregnb'), sg.Text("", size=(0, 1), key='last_usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcoregnb'), sg.Text("", size=(0, 1), key='last_passcore')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dircoregnb'), sg.Text("", size=(0, 1), key='last_dircore')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface40():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation_interface31.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Core data:', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcore'), sg.Text("", size=(0, 1), key='last_ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercore'), sg.Text("", size=(0, 1), key='last_usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcore',password_char='*'), sg.Text("", size=(0, 1), key='last_passcore')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dircore'), sg.Text("", size=(0, 1), key='last_dircore')],
        [sg.Text('gNB data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipgnb'), sg.Text("", size=(0, 1), key='last_ipgnb')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usergnb'), sg.Text("", size=(0, 1), key='last_usergnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passgnb',password_char='*'), sg.Text("", size=(0, 1), key='last_passgnb')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirgnb'), sg.Text("", size=(0, 1), key='last_dirgnb')],
        [sg.Text('OAI UE data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipoaiue'), sg.Text("", size=(0, 1), key='last_ipoaiue')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='useroaiue'), sg.Text("", size=(0, 1), key='last_useroaiue')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passoaiue',password_char='*'), sg.Text("", size=(0, 1), key='last_passoaiue')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirue'), sg.Text("", size=(0, 1), key='last_diroaiue')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 750))


def interface41():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/setupinformation_interface32.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Core data:', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipcore'), sg.Text("", size=(0, 1), key='last_ipcore')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercore'), sg.Text("", size=(0, 1), key='last_usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcore',password_char='*'), sg.Text("", size=(0, 1), key='last_passcore')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dircore'), sg.Text("", size=(0, 1), key='last_dircore')],
        [sg.Text('gNB data:')],
        [sg.Text('IP: ', size=(5, 1)), sg.Input('', key='ipgnb'), sg.Text("", size=(0, 1), key='last_ipgnb')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usergnb'), sg.Text("", size=(0, 1), key='last_usergnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passgnb',password_char='*'), sg.Text("", size=(0, 1), key='last_passgnb')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirgnb'), sg.Text("", size=(0, 1), key='last_dirgnb')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface42():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/imagens_interface42.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Information for gNB', size=(56, 1)),sg.Text('Previous information:')],
        [sg.Text('Name: ', size=(9, 1)), sg.Input('', key='namegnb'), sg.Text("", size=(0, 1), key='last_namegnb')],
        [sg.Text('Band: ', size=(9, 1)), sg.Input('', key='bandgnb'), sg.Text("", size=(0, 1), key='last_bandgnb')],
        [sg.Text('SCS: ', size=(9, 1)), sg.Input('', key='scsgnb'), sg.Text("", size=(0, 1), key='last_scsgnb')],
        [sg.Text('Bandwitdh: ', size=(9, 1)), sg.Input('', key='bwgnb'), sg.Text("", size=(0, 1), key='last_bwgnb')],
        [sg.Text('mcc: ', size=(9, 1)), sg.Input('', key='mccgnb'), sg.Text("", size=(0, 1), key='last_mccgnb')],
        [sg.Text('mnc: ', size=(9, 1)), sg.Input('', key='mncgnb'), sg.Text("", size=(0, 1), key='last_mncgnb')],
        [sg.Text('sst: ', size=(9, 1)), sg.Input('', key='sstgnb'), sg.Text("", size=(0, 1), key='last_sstgnb')],
        [sg.Text('sd: ', size=(9, 1)), sg.Input('', key='sdgnb'), sg.Text("", size=(0, 1), key='last_sdgnb')],
        [sg.Text('ID: ', size=(9, 1)), sg.Input('', key='idgnb'), sg.Text("", size=(0, 1), key='last_idgnb')],
        [sg.Text('Information for UE')],
        [sg.Text('IMSI: ', size=(9, 1)), sg.Input('', key='imsi'), sg.Text("", size=(0, 1), key='last_imsi')],
        [sg.Text('KEY: ', size=(9, 1)), sg.Input('', key='key'), sg.Text("", size=(0, 1), key='last_key')],
        [sg.Text('OPC: ', size=(9, 1)), sg.Input('', key='opc'), sg.Text("", size=(0, 1), key='last_opc')],
        [sg.Text('DNN: ', size=(9, 1)), sg.Input('', key='dnn'), sg.Text("", size=(0, 1), key='last_dnn')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface43():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface52_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Information for gNB', size=(56, 1)),sg.Text('Previous information:')],
        [sg.Text('Name: ', size=(9, 1)), sg.Input('', key='namegnb'), sg.Text("", size=(0, 1), key='last_namegnb')],
        [sg.Text('Band: ', size=(9, 1)), sg.Input('', key='bandgnb'), sg.Text("", size=(0, 1), key='last_bandgnb')],
        [sg.Text('SCS: ', size=(9, 1)), sg.Input('', key='scsgnb'), sg.Text("", size=(0, 1), key='last_scsgnb')],
        [sg.Text('Bandwitdh: ', size=(9, 1)), sg.Input('', key='bwgnb'), sg.Text("", size=(0, 1), key='last_bwgnb')],
        [sg.Text('mcc: ', size=(9, 1)), sg.Input('', key='mccgnb'), sg.Text("", size=(0, 1), key='last_mccgnb')],
        [sg.Text('mnc: ', size=(9, 1)), sg.Input('', key='mncgnb'), sg.Text("", size=(0, 1), key='last_mncgnb')],
        [sg.Text('sst: ', size=(9, 1)), sg.Input('', key='sstgnb'), sg.Text("", size=(0, 1), key='last_sstgnb')],
        [sg.Text('sd: ', size=(9, 1)), sg.Input('', key='sdgnb'), sg.Text("", size=(0, 1), key='last_sdgnb')],
        [sg.Text('ID: ', size=(9, 1)), sg.Input('', key='idgnb'), sg.Text("", size=(0, 1), key='last_idgnb')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface44():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/sim_card_programming.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Program the SIM Card')],
        [sg.Text(' ')],
        [sg.Text('Data for SIM Card:')],
        [sg.Text('ADM: ', size=(5, 1)), sg.Input('', key='adm')],
        [sg.Text('imsi: ', size=(5, 1)), sg.Input('', key='imsi')],
        [sg.Text('key: ', size=(5, 1)), sg.Input('', key='key')],
        [sg.Text('opc: ', size=(5, 1)), sg.Input('', key='opc')],
        [sg.Text('spn: ', size=(5, 1)), sg.Input('', key='spn')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 620))


def interface45():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface35_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('Identifies the NIC of each machine:')],
        [sg.Text('NIC Core:')],
        [sg.Combo(nic_core, key='niccore')],
        [sg.Text('NIC gNB:')],
        [sg.Combo(nic_gnb, key='nicgnb')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface46():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    layout = [
        [sg.Text('Preparing Setup....')],
        [sg.Text('Preparing configuration files and connecting setup...')],
        [sg.Text('Setup type: '), sg.Text("", size=(0, 1), key='setuptype')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Column([[gif]], justification='center')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface47():
    sg.theme('Reddit')
    my_img = ' '
    if show1:
        my_img = sg.Image(filename='Images/SetupOn_AllinOne_OAIUE.png', key='_CAMIMAGE_')
    if show2:
        my_img = sg.Image(filename='Images/SetupOn_AllinOne_COTSUE.png', key='_CAMIMAGE_')
    if show3:
        my_img = sg.Image(filename='Images/SetupOn_OAIUE.png', key='_CAMIMAGE_')
    if show4:
        my_img = sg.Image(filename='Images/SetupOn_COTSUE.png', key='_CAMIMAGE_')
    if show5:
        my_img = sg.Image(filename='Images/SetupOn_AllinOne_OAIUE.png', key='_CAMIMAGE_')
    if show6:
        my_img = sg.Image(filename='Images/SetupOn_AllinOne_COTSUE.png', key='_CAMIMAGE_')
    if show7:
        my_img = sg.Image(filename='Images/SetupOn_OAIUE.png', key='_CAMIMAGE_')
    if show8:
        my_img = sg.Image(filename='Images/SetupOn_COTSUE.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Setup ON')],
        [sg.Text('At this moment all network elements have already started, being ')],
        [sg.Text('able to verify this information in the core amf logs. ')],
        [sg.Text(' ')],
        [sg.Text('You can use the network at any time. ')],
        [sg.Text(' ')],
        [sg.Text('If you want to disconnect this network, just click on the "Finished"  ')],
        [sg.Text('button or close the window. ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Finished')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))

#################################################
#    Interface para testes (configuração)       #
#################################################


def interface48():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface48_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text('Select one of the two options')],
        [sg.Text('Note: If it is the first time to carry out the tests, it is necessary to configure the platform')],
        [sg.Radio('Configuration', "RADIO1", key='config')],
        [sg.Radio('Tests', "RADIO1", key='tests')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 590))


def interface49():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface49_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('IP address of each machine: ', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('Core: ', size=(5, 1)), sg.Input('', key='ipcore'), sg.Text("", size=(0, 1), key='last_ipcore')],
        [sg.Text('gNB: ', size=(5, 1)), sg.Input('', key='ipgnb'), sg.Text("", size=(0, 1), key='last_ipgnb')],
        [sg.Text('UE: ', size=(5, 1)), sg.Input('', key='ipue'), sg.Text("", size=(0, 1), key='last_ipoaiue')],
        [sg.Text(' ')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 620))


def interface50():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface50_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('Authentication data of the machines on which the OAI is running')],
        [sg.Text('Core: ', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usercore'), sg.Text("", size=(0, 1), key='last_usercore')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passcore', password_char='*'), sg.Text("", size=(0, 1), key='last_passcore')],
        [sg.Text('gNB: ')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='usergnb'), sg.Text("", size=(0, 1), key='last_usergnb')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passgnb', password_char='*'), sg.Text("", size=(0, 1), key='last_passgnb')],
        [sg.Text('UE: ')],
        [sg.Text('User: ', size=(5, 1)), sg.Input('', key='userue'), sg.Text("", size=(0, 1), key='last_useroaiue')],
        [sg.Text('Pass: ', size=(5, 1)), sg.Input('', key='passue', password_char='*'), sg.Text("", size=(0, 1), key='last_passoaiue')],
        [sg.Text(' ')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 620))


def interface51():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface51_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('Directories where the OAI is located on each machine')],
        [sg.Text('Core: ', size=(52, 1)),sg.Text('Previous information:')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dircore'), sg.Text("", size=(0, 1), key='last_dircore')],
        [sg.Text('gNB: ')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirgnb'), sg.Text("", size=(0, 1), key='last_dirgnb')],
        [sg.Text('UE: ')],
        [sg.Text('Path: ', size=(5, 1)), sg.Input('', key='dirue'), sg.Text("", size=(0, 1), key='last_diroaiue')],
        [sg.Text(' ')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 610))


def interface52():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface52_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text('gNB settings: ', size=(72, 1)),sg.Text('Previous information:')],
        [sg.Text('Name: ', size=(25, 1)), sg.Input('', key='gnbname'), sg.Text("", size=(0, 1), key='last_gnbname')],
        [sg.Text('ID: : ', size=(25, 1)), sg.Input('', key='gnbid'), sg.Text("", size=(0, 1), key='last_gnbid')],
        [sg.Text('MCC: : ', size=(25, 1)), sg.Input('', key='mcc'), sg.Text("", size=(0, 1), key='last_mcc')],
        [sg.Text('MNC: ', size=(25, 1)), sg.Input('', key='mnc'), sg.Text("", size=(0, 1), key='last_mnc')],
        [sg.Text('SST : ', size=(25, 1)), sg.Input('', key='sst'), sg.Text("", size=(0, 1), key='last_sst')],
        [sg.Text('SD : ', size=(25, 1)), sg.Input('', key='sd'), sg.Text("", size=(0, 1), key='last_sd')],
        [sg.Text('AMF ip: ', size=(25, 1)), sg.Input('', key='amfip'), sg.Text("", size=(0, 1), key='last_amfip')],
        [sg.Text('gNB interface name for AMF: ', size=(25, 1)), sg.Input('', key='gnbnameamf'), sg.Text("", size=(0, 1), key='last_gnbnameamf')],
        [sg.Text('gNB ip for AMF: ', size=(25, 1)), sg.Input('', key='gnbipamf'), sg.Text("", size=(0, 1), key='last_gnbipamf')],
        [sg.Text('gNB interface name for NGU: ', size=(25, 1)), sg.Input('', key='gnbnamegnu'), sg.Text("", size=(0, 1), key='last_gnbnamegnu')],
        [sg.Text('gNB ip for NGU: ', size=(25, 1)), sg.Input('', key='gnbipngu'), sg.Text("", size=(0, 1), key='last_gnbipngu')],
        [sg.Text(' ')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface53():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface53_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='center')],
        [sg.Text(' ')],
        [sg.Text('Finish settings')],
        [sg.Text(' ')],
        [sg.Text('Click on the "Back" button if you want to change any settings previously introduced')],
        [sg.Text(' ')],
        [sg.Text('Click on the "Save" button if you want to save the settings')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Save')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 610))


# tests
def interface54():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface54_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('Choose the type of test:')],
        [sg.Radio('Check if Core is on', "RADIO1", key='coreon')],
        [sg.Radio('Check if UE is connected', "RADIO1", key='ueon')],
        [sg.Radio('End-to-end test', "RADIO1", key='uetoext')],
        [sg.Radio('Test using iperf', "RADIO1", key='iperftest')],
        [sg.Text(' ')],
        [sg.Text('Save:')],
        [sg.Checkbox('Document with gNB logs(.pdf)', key='relgnb')],
        [sg.Checkbox('Document with UE logs(.pdf)', key='relue')],
        [sg.Checkbox('Report test(.pdf)', key='reltest')],
        [sg.Text('Local: ', size=(5, 1)), sg.Input('', key='diretorio')],
        #[sg.Text('Local: ', size=(5, 1)), sg.Input(key='diretorio'), sg.FileBrowse('Select')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Next')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 660))


def interface55():
    sg.theme('Reddit')
    gif = sg.Image(sg.DEFAULT_BASE64_LOADING_GIF, key='-GIF-')
    if coreon:
        my_img = sg.Image(filename='Images/interface55_coretest_image.png', key='_CAMIMAGE_')
        layout = [
            [sg.Column([[my_img]], justification='left')],
            [sg.Text('Running the Core Test...')],
            [sg.Column([[gif]], justification='center')]
        ]
    if ueon:
        my_img = sg.Image(filename='Images/interface55_uetest_image.png', key='_CAMIMAGE_')
        layout = [
            [sg.Column([[my_img]], justification='left')],
            [sg.Text('Running the UE Test...')],
            [sg.Column([[gif]], justification='center')]
        ]
    if ueconnect:
        my_img = sg.Image(filename='Images/interface55_ueconncetion_image.png', key='_CAMIMAGE_')
        layout = [
            [sg.Column([[my_img]], justification='left')],
            [sg.Text('Running the UE internet connection test...')],
            [sg.Column([[gif]], justification='center')]
        ]
    if iperf:
        my_img = sg.Image(filename='Images/interface55_iperftest_image.png', key='_CAMIMAGE_')
        layout = [
            [sg.Column([[my_img]], justification='left')],
            [sg.Text('Running the test with Iperf...')],
            [sg.Column([[gif]], justification='center')]
        ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 650))


def interface56():
    sg.theme('Reddit')
    if result_test_coreon:
        my_img = sg.Image(filename='Images/interface56_on.png', key='_CAMIMAGE_')
        layout = [
            [sg.Text('Results:')],
            [sg.Text('Successful test, Core is online')],
            [sg.Text(' ')],
            [sg.Text('Test conditions:')],
            [sg.Text('Core Machine with IP: '), sg.Text("", size=(0, 1), key='ipcoreout')],
            [sg.Text('AMF IP address:'), sg.Text("", size=(0, 1), key='ipamfout')],
            [sg.Text('Test machine IP'), sg.Text("", size=(0, 1), key='ipranout')],
            [sg.Text(' ')],
            [sg.Text(' ')],
            [sg.Text(' ')],
            [sg.Column([[my_img]], justification='center')],
            [sg.Text(' ')],
            [sg.Button('Close')]
        ]

    if not result_test_coreon:
        my_img = sg.Image(filename='Images/interface56_erro.png', key='_CAMIMAGE_')
        layout = [
            [sg.Text('Results: Core is offline')],
            [sg.Text(' ')],
            [sg.Text('Possible causes of the problem:')],
            [sg.Text('-> Core is not running')],
            [sg.Text('-> Core configuration problems')],
            [sg.Text('-> Machine where the Core is running is disconnected')],
            [sg.Text(' ')],
            [sg.Text('What to do?')],
            [sg.Text('-> Check if the Core is running, if not, run the command to run the Core')],
            [sg.Text('-> Run the commands " sudo sysctl net.ipv4.conf.all.forwarding=1 " and '
                     '" sudo iptables -P FORWARD ACCEPT "')],
            [sg.Text('-> Check if there is a problem with the network card of the machine running the Core')],
            [sg.Text('-> Check if the address of the machine on which the Core runs is correct')],
            [sg.Text('-> Check if the network cable is connected')],
            [sg.Text(' ')],
            [sg.Column([[my_img]], justification='center')],
            [sg.Text(' ')],
            [sg.Text(' ')],
            [sg.Button('Close')]
        ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 620))


def interface57():
    sg.theme('Reddit')
    my_img = sg.Image(filename='Images/interface57_image.png', key='_CAMIMAGE_')
    layout = [
        [sg.Column([[my_img]], justification='left')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text('Information about the bandwidth, scs and bandwidth of the test:')],
        [sg.Text(' ')],
        [sg.Text('Band: ')],
        [sg.Combo(band_interface(), key='band')],
        [sg.Text('SCS: ')],
        [sg.Combo(scs_interface(), key='scs')],
        [sg.Text('BW: ')],
        [sg.Combo(bw_interface(), key='bw')],
        [sg.Text(' ')],
        [sg.Checkbox('Use previous information', key='oldinfo')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Text(' ')],
        [sg.Button('Back'), sg.Button('Start')]
    ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 600))


def interface58():
    sg.theme('Reddit')
    if result_test_ueon:
        my_img = sg.Image(filename='Images/interface58_on.png', key='_CAMIMAGE_')
        layout = [
            [sg.Text('Results: EU is connected')],
            [sg.Text(' ')],
            [sg.Text('Test conditions:')],
            [sg.Text('Core:')],
            [sg.Text('Machine IP: '), sg.Text("", size=(0, 1), key='ipcoreout')],
            [sg.Text('AMF IP address:'), sg.Text("", size=(0, 1), key='ipamfout')],
            [sg.Text(' ')],
            [sg.Text('gNB:')],
            [sg.Text('Machine IP:'), sg.Text("", size=(0, 1), key='ipranout')],
            [sg.Text('Band:'), sg.Text("", size=(0, 1), key='bandout')],
            [sg.Text('Subcarrier Spacing:'), sg.Text("", size=(0, 1), key='scsout')],
            [sg.Text('BandWidth:'), sg.Text("", size=(0, 1), key='bwout')],
            [sg.Text(' ')],
            [sg.Text('UE:')],
            [sg.Text('Machine IP:'), sg.Text("", size=(0, 1), key='ipueoutmachine')],
            [sg.Text('UE IP:'), sg.Text("", size=(0, 1), key='ipueout')],
            [sg.Column([[my_img]], justification='center')],
            [sg.Button('Close')]
        ]

    if not result_test_ueon:
        my_img = sg.Image(filename='Images/interface58_erro.png', key='_CAMIMAGE_')
        layout = [
            [sg.Text('Result: UE is not connected')],
            [sg.Text('Possible causes:')],
            [sg.Text('-> Core or gNB is offline')],
            [sg.Text('-> Core configuration problems')],
            [sg.Text('-> Problems in the connection between Core and gNB')],
            [sg.Text('-> Missing/Incorrect route configuration on the machine running gNB')],
            [sg.Text('-> mcc, mnc, sst, sd, imsi, key, opc, and dnn values not matching Core, gNB, and UE')],
            [sg.Text(' ')],
            [sg.Text('What to do?')],
            [sg.Text('-> Check if gNB and Core are online')],
            [sg.Text('-> Check Core, gNB and UE configuration')],
            [sg.Text('-> Check if the machines that run Core and gNB have all the correct configurations')],
            [sg.Text(' ')],
            [sg.Text('Suggestions:')],
            [sg.Text('-> Repeat this test, saving the logs if they have not done so and analyzing '
                     'the information contained')],
            [sg.Text('-> Run a verification test on the Core state in order to exclude some '
                     'previously mentioned error sources')],
            [sg.Column([[my_img]], justification='center')],
            [sg.Button('Close')]
        ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 700))


def interface59():
    sg.theme('Reddit')
    if result_test_ueconnect:
        my_img = sg.Image(filename='Images/interface59_on.png', key='_CAMIMAGE_')
        layout = [
            [sg.Text('Results: Test successful')],
            [sg.Text(' ')],
            [sg.Text('Test conditions:')],
            [sg.Text('Core:')],
            [sg.Text('Machine IP: '), sg.Text("", size=(0, 1), key='ipcoreout')],
            [sg.Text('AMF IP address:'), sg.Text("", size=(0, 1), key='ipamfout')],
            [sg.Text(' ')],
            [sg.Text('gNB:')],
            [sg.Text('Machine IP:'), sg.Text("", size=(0, 1), key='ipranout')],
            [sg.Text('Band:'), sg.Text("", size=(0, 1), key='bandout')],
            [sg.Text('Subcarrier Spacing:'), sg.Text("", size=(0, 1), key='scsout')],
            [sg.Text('BandWidth:'), sg.Text("", size=(0, 1), key='bwout')],
            [sg.Text(' ')],
            [sg.Text('UE:')],
            [sg.Text('Machine IP:'), sg.Text("", size=(0, 1), key='ipueoutmachine')],
            [sg.Text('UE IP:'), sg.Text("", size=(0, 1), key='ipueout')],
            [sg.Column([[my_img]], justification='center')],
            [sg.Button('Close')]
        ]

    if not result_test_ueconnect:
        my_img = sg.Image(filename='Images/interface59_erro.png', key='_CAMIMAGE_')
        layout = [
            [sg.Text('Result: Test failed')],
            [sg.Text('Possible causes:')],
            [sg.Text('-> Core offline')],
            [sg.Text('-> No connection between gNB and Core')],
            [sg.Text('-> Wrong UE, gNB, or Core settings')],
            [sg.Text('-> mcc, mnc, sst, sd, imsi, key, opc, and dnn values not matching Core, gNB, and UE')],
            [sg.Text(' ')],
            [sg.Text('What to do?')],
            [sg.Text('-> Check if machines running UE, gNB, and Core are turned on')],
            [sg.Text('-> Check if Core, gNB, and UE are running (analyze logs)')],
            [sg.Text('-> Check if the network cables are connected')],
            [sg.Text('-> Check if the IP addresses are correct')],
            [sg.Text(' ')],
            [sg.Text('Suggestion:')],
            [sg.Text(
                '-> Run the test that checks if the UE connects, if the test is successful, check the ')],
            [sg.Text('docker-compose file and check if the IP address of the interface that connects to'
                     ' the outside agrees with that of the NIC')],
            [sg.Column([[my_img]], justification='center')],
            [sg.Button('Close')]
        ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 800))


def interface60():
    sg.theme('Reddit')
    if result_test_iperf:
        my_img = sg.Image(filename='Images/interface60_on.png', key='_CAMIMAGE_')
        layout = [
            [sg.Text('Result: Test successful')],
            [sg.Text('Test conditions:')],
            [sg.Text('Core:')],
            [sg.Text('Machine IP: '), sg.Text("", size=(0, 1), key='ipcoreout')],
            [sg.Text('AMF IP address:'), sg.Text("", size=(0, 1), key='ipamfout')],
            [sg.Text('gNB:')],
            [sg.Text('Machine IP: '), sg.Text("", size=(0, 1), key='ipranout')],
            [sg.Text('Band:'), sg.Text("", size=(0, 1), key='bandout')],
            [sg.Text('Subcarrier Spacing:'), sg.Text("", size=(0, 1), key='scsout')],
            [sg.Text('BandWidth:'), sg.Text("", size=(0, 1), key='bwout')],
            [sg.Text('UE:')],
            [sg.Text('Machine IP: '), sg.Text("", size=(0, 1), key='ipueoutmachine')],
            [sg.Text('UE IP: '), sg.Text("", size=(0, 1), key='ipueout')],
            [sg.Text('Iperf data:')],
            [sg.Text('Test duration: 10s')],
            [sg.Text('Traffic:'), sg.Text("", size=(0, 1), key='tipotrafego')],
            [sg.Text('Transferred data:'), sg.Text("", size=(0, 1), key='quantidade')],
            [sg.Text('Packages lost:'), sg.Text("", size=(0, 1), key='perdidos')],
            [sg.Text('Total packages:'), sg.Text("", size=(0, 1), key='total')],
            [sg.Text('Percentage of lost packets:'), sg.Text("", size=(0, 1), key='percentagem')],
            [sg.Column([[my_img]], justification='center')],
            [sg.Button('Close')]
        ]

    if not result_test_iperf:
        my_img = sg.Image(filename='Images/interface60_erro.png', key='_CAMIMAGE_')
        layout = [
            [sg.Text('Result: Test failed')],
            [sg.Text('Possible causes:')],
            [sg.Text('-> Machines running gNB or Core are not properly configured')],
            [sg.Text('-> Wrong configuration files')],
            [sg.Text('-> Core and gNB are not connected')],
            [sg.Text(' ')],
            [sg.Text('What to do?')],
            [sg.Text('-> Check if network elements turn on')],
            [sg.Text('-> Check the settings of the different machines')],
            [sg.Text('-> Check configuration files')],
            [sg.Text(' ')],
            [sg.Text('Suggestion:')],
            [sg.Text('-> Run the Core verification and UE connection verification tests,')],
            [sg.Text('in order to more easily identify the source of the problem')],
            [sg.Column([[my_img]], justification='center')],
            [sg.Button('Close')]
        ]
    return sg.Window('OAI Optimized Platform', layout=layout, finalize=True, size=(850, 750))


#################################################
#                Funções Gerais                 #
#################################################
def stop_core():
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipcore, username=usercore, password=passcore)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('cd /'+str(dircore)+'/oai-cn5g-fed/docker-compose; sudo -S python3 '
                                                            'core-network.py --type stop-basic;')
    ssh_stdin.write(passcore+'\n')
    ssh_stdin.flush()
    ssh_stdin.close()


def stop_gnb():
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipgnb, username=usergnb, password=passgnb)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S pkill nr-softmodem')
    ssh_stdin.write(passgnb+'\n')
    ssh_stdin.flush()

    print('Stop gNB')

    time.sleep(2)

    if relgnb:
        if os.path.isfile('logs_gNB.txt'):
            os.remove('logs_gNB.txt')

        dirlocal = subprocess.getoutput('pwd')
        sftp = client.open_sftp()
        localpath = dirlocal + '/logs_gNB.txt'
        remotepath = '/' + str(dirgnb) + '/openairinterface5g/cmake_targets/ran_build/build/logs_gNB.txt'
        sftp.get(remotepath, localpath)
        sftp.close()

        print('Entrou para fazer o relatorio do gNB')
        if os.path.isfile('logs_gNB.txt'):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=8)
            file = open("logs_gNB.txt", "r")

            for i in file:
                pdf.cell(10, 5, txt=i, ln=1, align="l")

            pdf.output("Logs_gNB.pdf")

    ssh_stdin.close()


def stop_oaiue():
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S pkill nr-uesoftmodem')
    ssh_stdin.write(passoaiue + '\n')
    ssh_stdin.flush()
    ######
    if relue:
        if os.path.isfile('logs_oaiue.txt'):
            os.remove('logs_oaiue.txt')

        dirlocal = subprocess.getoutput('pwd')
        sftp = client.open_sftp()
        localpath = dirlocal + '/logs_oaiue.txt'
        remotepath = '/' + str(dirue) + '/openairinterface5g/cmake_targets/ran_build/build/logs_oaiue.txt'
        sftp.get(remotepath, localpath)
        sftp.close()

        print('Entrou para fazer o relatorio do OAIUE')
        if os.path.isfile('logs_oaiue.txt'):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=8)
            file = open("logs_oaiue.txt", "r")

            for i in file:
                pdf.cell(10, 5, txt=i, ln=1, align="l")

            pdf.output("Logs_oaiue.pdf")

    ssh_stdin.close()


def limpar():
    if os.path.isfile('finalizadocore.txt'):
        os.remove('finalizadocore.txt')
    if os.path.isfile('finalizadognb.txt'):
        os.remove('finalizadognb.txt')
    if os.path.isfile('finalizadooaiue.txt'):
        os.remove('finalizadooaiue.txt')
    if os.path.isfile('docker-compose-basic-nrf.yaml'):
        os.remove('docker-compose-basic-nrf.yaml')
    if os.path.isfile('dns.txt'):
        os.remove('dns.txt')
    if os.path.isfile('ocupadas.txt'):
        os.remove('ocupadas.txt')
    if os.path.isfile('todas.txt'):
        os.remove('todas.txt')
    if os.path.isfile('configure.conf'):
        os.remove('configure.conf')
    if os.path.isfile('ue.conf'):
        os.remove('ue.conf')
    if os.path.isfile('feito_gnb.txt'):
        os.remove('feito_gnb.txt')
    if os.path.isfile('feito_oaiue.txt'):
        os.remove('feito_oaiue.txt')
    if os.path.isfile('oai_db.sql'):
        os.remove('oai_db.sql')
    if os.path.isfile('coreon_result.txt'):
        os.remove('coreon_result.txt')
    if os.path.isfile('ueon_result.txt'):
        os.remove('ueon_result.txt')
    if os.path.isfile('ueconnect_result.txt'):
        os.remove('ueconnect_result.txt')
    if os.path.isfile('ip_test.txt'):
        os.remove('ip_test.txt')
    if os.path.isfile('result_iperf.txt'):
        os.remove('result_iperf.txt')
    if os.path.isfile('stop.txt'):
        os.remove('stop.txt')
    if os.path.isfile('logs_gNB.txt'):
        os.remove('logs_gNB.txt')
    if os.path.isfile('dest1.docx'):
        os.remove('dest1.docx')
    if os.path.isfile('logs_oaiue.txt'):
        os.remove('logs_oaiue.txt')
    if os.path.isfile('stop_ue.txt'):
        os.remove('stop_ue.txt')

#################################################
#                     Start                     #
#################################################


def run_core():
    time.sleep(20)

    dirlocal = subprocess.getoutput('pwd')

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipcore, username=usercore, password=passcore)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S ip addr flush '+str(niccore)+'; sudo -S ip addr '
                                                                    'add 192.168.1.16/24 dev '+str(niccore)+';')
    ssh_stdin.write(passcore+'\n')
    ssh_stdin.flush()
    time.sleep(1)

    sftp = client.open_sftp()
    localpath = dirlocal + '/oai_db.sql'
    remotepath = '/'+str(dircore)+'/oai-cn5g-fed/docker-compose/database/oai_db.sql'
    print(remotepath)
    sftp.put(localpath, remotepath)
    sftp.close()

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S chmod 644 /'+str(dircore)+'/oai-cn5g-fed'
                                                            '/docker-compose/database/oai_db.sql;')

    ssh_stdin.write(passcore + '\n')
    ssh_stdin.flush()
    time.sleep(1)

    sftp = client.open_sftp()
    localpath = dirlocal + '/docker-compose-basic-nrf.yaml'
    remotepath = '/'+str(dircore)+'/oai-cn5g-fed/docker-compose/docker-compose-basic-nrf.yaml'
    print(remotepath)
    sftp.put(localpath, remotepath)
    sftp.close()


    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S sysctl net.ipv4.conf.all.forwarding=1;'
                                                            ' sudo -S iptables -P FORWARD ACCEPT')
    ssh_stdin.write(passcore+'\n')
    ssh_stdin.flush()

    time.sleep(5)

    cmd_exec_core = 'cd /'+str(dircore)+'/oai-cn5g-fed/docker-compose/; python3 core-network.py --type ' \
                                             'start-basic --scenario 1; cd; feito > feito.txt'
    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command(cmd_exec_core)
    ssh_stdin.flush()

    data_core = ssh_stdout.readlines()  # resolveu com isto
    ssh_stdin.close()


def run_gnb():
    time.sleep(80)

    print('Run gNB')

    dirlocal = subprocess.getoutput('pwd')

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipgnb, username=usergnb, password=passgnb)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S ip addr flush '+str(nicgnb)+'; sudo -S ip addr add 192.168.1.15/24 dev '+str(nicgnb)+';')
    ssh_stdin.write(passgnb+'\n')
    ssh_stdin.flush()
    time.sleep(1)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S ip route add 192.168.70.128/26 via 192.168.1.16 dev '+str(nicgnb)+';')
    ssh_stdin.write(passgnb+'\n')
    ssh_stdin.flush()

    time.sleep(5)

    os.system('echo feito > feito_gnb.txt')
    # Problema aqui:
    sftp = client.open_sftp()
    localpath = dirlocal + '/configure.conf'
    remotepath = '/'+str(dirgnb)+'/openairinterface5g/targets/PROJECTS/GENERIC-NR-5GC/CONF/configure.conf'
    sftp.put(localpath, remotepath)
    sftp.close()

    time.sleep(5) # dar tempo para atualizar

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('cd /' + str(dirgnb) + '/openairinterface5g/cmake_targets/ran_build/build; sudo -S ./nr-softmodem -O /' + str(dirgnb) + '/openairinterface5g/targets/PROJECTS/GENERIC-NR-5GC/CONF/configure.conf --sa --continuous-tx -E;')
    ##################################################################################################################
    ssh_stdin.write(passgnb+'\n')
    ssh_stdin.flush()

    data_gnb = ssh_stdout.readlines()  # resolveu com isto
    ssh_stdin.close()


def run_oaiue():

    dirlocal = subprocess.getoutput('pwd')

    time.sleep(130)
    print('Run UE')
    os.system('echo feito > feito_oaiue.txt')

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)

    sftp = client.open_sftp()
    localpath = dirlocal + '/ue.conf'
    remotepath = '/'+str(dirue)+'/openairinterface5g/targets/PROJECTS/GENERIC-NR-5GC/CONF/ue.conf'
    sftp.put(localpath, remotepath)
    sftp.close()

    time.sleep(5)

    os.system('echo feito > feito_oaiue.txt')

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command(
        'cd /'+str(dirue)+'/openairinterface5g/cmake_targets/ran_build/build; sudo -S ./nr-uesoftmodem -r 106 --numerology 1 --band 78 -C 3619200000 --ue-fo-compensation --sa --nokrnmod -O /'+str(dirue)+'/openairinterface5g/targets/PROJECTS/GENERIC-NR-5GC/CONF/ue.conf -E;')
    ssh_stdin.write(passoaiue + '\n')
    ssh_stdin.flush()

    data_oaiue = ssh_stdout.readlines()  # resolveu com isto
    ssh_stdin.close()

#################################################
#          Configurações de instalação          #
#################################################


def corebash():
    if os.path.isfile('Templete/installcorefirst_templete.sh'):
        with open('Templete/installcorefirst_templete.sh', 'r') as filecore:
            data = filecore.read()
    filecore.close()

    data = data.replace('xxxxx', usercore)

    with open('Shell/installcorefirst.sh', 'w') as filecore:
        filecore.write(data)

    if os.path.isfile('Templete/installcoresecond_templete.sh'):
        with open('Templete/installcoresecond_templete.sh', 'r') as filecore:
            data = filecore.read()
    filecore.close()

    data = data.replace('yyyyy', passcore)

    with open('Shell/installcoresecond.sh', 'w') as filecore:
        filecore.write(data)


def gnbbash():
    if os.path.isfile('Templete/installgnb_templete.sh'):
        with open('Templete/installgnb_templete.sh', 'r') as filegnb:
            data = filegnb.read()
    filegnb.close()

    data = data.replace('yyyyy', passgnb)

    with open('Shell/installgnb.sh', 'w') as filegnb:
        filegnb.write(data)

#################################################
#                  Instaladores                 #
#################################################


def do_install_core():
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipcore, username=usercore, password=passcore)

    dirlocal = subprocess.getoutput('pwd')

    sftp = client.open_sftp()
    localpath = dirlocal+'/Shell/clean.sh'
    remotepath = '/home/'+usercore+'/clean.sh'
    sftp.put(localpath, remotepath)
    sftp.close()

    sftp = client.open_sftp()
    localpath = dirlocal+'/Shell/installcorefirst.sh'
    remotepath = '/home/'+usercore+'/installcorefirst.sh'
    sftp.put(localpath, remotepath)
    sftp.close()

    sftp = client.open_sftp()
    localpath = dirlocal+'/Shell/installcoresecond.sh'
    remotepath = '/home/'+usercore+'/installcoresecond.sh'
    sftp.put(localpath, remotepath)
    sftp.close()

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command(
        'sudo -S chmod +x installcorefirst.sh; sudo -S chmod +x installcoresecond.sh; sudo -S chmod +x clean.sh;')
    ssh_stdin.write(passcore+'\n')
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    print(data)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S ./installcorefirst.sh;')
    ssh_stdin.write(passcore+'\n')
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    print(data)

    time.sleep(200)

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipcore, username=usercore, password=passcore)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('./installcoresecond.sh;')
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    print(data)

    sftp = client.open_sftp()
    localpath = dirlocal+'/finalizadocore.txt'
    remotepath = '/home/'+usercore+'/finalizadocore.txt'
    sftp.get(remotepath, localpath)
    sftp.close()


def do_install_gnb():
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipgnb, username=usergnb, password=passgnb)

    dirlocal = subprocess.getoutput('pwd')

    sftp = client.open_sftp()
    localpath = dirlocal+'/Shell/cleangnb.sh'
    remotepath = '/home/'+usergnb+'/cleangnb.sh'
    sftp.put(localpath, remotepath)
    sftp.close()

    sftp = client.open_sftp()
    localpath = dirlocal+'/Shell/installgnb.sh'
    remotepath = '/home/'+usergnb+'/installgnb.sh'
    sftp.put(localpath, remotepath)
    sftp.close()

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S chmod +x installgnb.sh;'
                                                            ' sudo -S chmod +x cleangnb.sh;')
    ssh_stdin.write(passgnb+'\n')
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    print(data)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('./installgnb.sh;')
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    print(data)

    sftp = client.open_sftp()
    localpath = dirlocal+'/finalizadognb.txt'
    remotepath = '/home/'+usergnb+'/finalizadognb.txt'
    sftp.get(remotepath, localpath)
    sftp.close()


def do_install_oaiue():
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)

    dirlocal = subprocess.getoutput('pwd')

    #sftp = client.open_sftp()
    #localpath = dirlocal+'/Shell/cleanoaiue.sh'
    #remotepath = '/home/'+useroaiue+'/cleanoaiue.sh'
    #sftp.put(localpath, remotepath)
    #sftp.close()

    sftp = client.open_sftp()
    localpath = dirlocal+'/Shell/installoaiue.sh'
    remotepath = '/home/'+useroaiue+'/installoaiue.sh'
    sftp.put(localpath, remotepath)
    sftp.close()

    #ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S chmod +x installoaiue.sh;'
    #                                                        ' sudo -S chmod +x cleanoaiue.sh; ./installoaiue.sh;')
    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S chmod +x installoaiue.sh;')

    ssh_stdin.write(passoaiue+'\n')
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    print(data)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('./installoaiue.sh;')
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    print(data)

    sftp = client.open_sftp()
    localpath = dirlocal+'/finalizadooaiue.txt'
    remotepath = '/home/'+useroaiue+'/finalizadooaiue.txt'
    sftp.get(remotepath, localpath)
    sftp.close()


def do_install_coregnb():
    do_install_core()
    do_install_gnb()


#################################################
#                 Configurações                 #
#################################################

def nic_information(ip, user, password): # Nic livres
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ip, username=user, password=password)

    dirlocal = subprocess.getoutput('pwd')

    sftp = client.open_sftp()
    localpath = dirlocal+'/Shell/interfaces.sh'
    remotepath = '/home/' + user + '/interfaces.sh'
    sftp.put(localpath, remotepath)
    sftp.close()

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('sudo -S chmod +x interfaces.sh;')
    ssh_stdin.write(password + '\n')
    ssh_stdin.flush()

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('./interfaces.sh')
    ssh_stdin.flush()

    sftp = client.open_sftp()
    localpath = dirlocal+'/ocupadas.txt'
    remotepath = '/home/' + user + '/ocupadas.txt'
    sftp.get(remotepath, localpath)
    sftp.close()

    if os.path.isfile('ocupadas.txt'):
        with open('ocupadas.txt', 'r') as file:
            ocupadas = file.read()
    file.close()

    sftp = client.open_sftp()
    localpath = dirlocal+'/todas.txt'
    remotepath = '/home/' + user + '/todas.txt'
    sftp.get(remotepath, localpath)
    sftp.close()

    placas = []

    todas = open('todas.txt', 'r')
    lines = todas.readlines()
    for line in lines:
        line = str(line)
        placas.append(line[0:len(line) - 1])

    return placas

###########################################
#           All in One - Fast             #
###########################################


def core_config_allinone(): # repetido do sleep core/gNB mas para ser rapido agora fica assim, pode ser para apagar
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipcore, username=usercore, password=passcore)

    dirlocal = subprocess.getoutput('pwd')

    cmd = 'netstat -ie | grep -B1 "' + ipcore + '" | head -n1 | awk ' + '\'' + '{print $1}' + '\'' + ';'
    print(cmd)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command(cmd)
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    data = str(data)
    interfacecore = data[2:len(data) - 5]
    print(interfacecore)

    cmd = 'nmcli device show ' + interfacecore + ' | grep IP4.DNS > dns.txt'
    client.exec_command(cmd)

    sftp = client.open_sftp()
    localpath = dirlocal + '/dns.txt'
    remotepath = '/home/' + usercore + '/dns.txt'
    sftp.get(remotepath, localpath)
    sftp.close()

    if os.path.isfile('dns.txt'):
        with open('dns.txt', 'r') as filecoreallinone:
            dns_data = filecoreallinone.read()
    filecoreallinone.close()

    dns1 = str(dns_data[40:53])
    dns2 = str(dns_data[94:len(dns_data)])

    print(dns1)
    print(dns2)

    if os.path.isfile('Templete/docker-compose-basic-nrf_templete_fast.yaml'):
        with open('Templete/docker-compose-basic-nrf_templete_fast.yaml', 'r') as filecoreallinone:
            docker_compose = filecoreallinone.read()
    filecoreallinone.close()

    docker_compose = docker_compose.replace('dns1', dns1)
    docker_compose = docker_compose.replace('dns2', dns2)

    with open('docker-compose-basic-nrf.yaml', 'w') as file:
        file.write(docker_compose)

    if os.path.isfile('Templete/oai_dbCore_gNB_fast.sql'):
        with open('Templete/oai_dbCore_gNB_fast.sql', 'r') as filecoreallinone:
            data_base = filecoreallinone.read()
    filecoreallinone.close()

    with open('oai_db.sql', 'w') as filecoreallinone:
        filecoreallinone.write(data_base)


def gnb_config_allinone():  # sem testar, possivelmente pode ser para apagar
    if os.path.isfile('Templete/ConfigureAllInOne_template_fast.conf'):
        with open('Templete/ConfigureAllInOne_template_fast.conf', 'r') as filegnballinone:
            gnb_data = filegnballinone.read()
    filegnballinone.close()

    with open('configure.conf', 'w') as filegnballinone:
        filegnballinone.write(gnb_data)


def oai_config_allinone():
    if os.path.isfile('Templete/ue_fast_templete.conf'):
        with open('Templete/ue_fast_templete.conf', 'r') as fileoaiallinone:
            oaiue_data = fileoaiallinone.read()

    fileoaiallinone.close()

    with open('ue.conf', 'w') as fileoaiallinone:
        fileoaiallinone.write(oaiue_data)


###########################################
#           Core e gNB - Fast             #
###########################################


def core_config():
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipcore, username=usercore, password=passcore)

    dirlocal = subprocess.getoutput('pwd')

    cmd = 'netstat -ie | grep -B1 "'+ipcore+'" | head -n1 | awk ' + '\'' + '{print $1}' + '\'' + ';'
    print(cmd)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command(cmd)
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    data = str(data)
    interfacecore = data[2:len(data) - 5]
    print(interfacecore)

    cmd = 'nmcli device show ' + interfacecore + ' | grep IP4.DNS > dns.txt'
    client.exec_command(cmd)

    sftp = client.open_sftp()
    localpath = dirlocal+'/dns.txt'
    remotepath = '/home/'+usercore+'/dns.txt'
    sftp.get(remotepath, localpath)
    sftp.close()

    if os.path.isfile('dns.txt'):
        with open('dns.txt', 'r') as filecoreconfig:
            dns_data = filecoreconfig.read()
    filecoreconfig.close()

    dns1 = str(dns_data[40:53])
    dns2 = str(dns_data[94:len(dns_data)])

    print(dns1)
    print(dns2)

    if os.path.isfile('Templete/docker-compose-basic-nrf_templete_fast.yaml'):
        with open('Templete/docker-compose-basic-nrf_templete_fast.yaml', 'r') as filecoreconfig:
            docker_compose = filecoreconfig.read()
    filecoreconfig.close()

    docker_compose = docker_compose.replace('dns1', dns1)
    docker_compose = docker_compose.replace('dns2', dns2)

    with open('docker-compose-basic-nrf.yaml', 'w') as filecoreconfig:
        filecoreconfig.write(docker_compose)

    if os.path.isfile('Templete/oai_dbCore_gNB_fast.sql'):
        with open('Templete/oai_dbCore_gNB_fast.sql', 'r') as filecoreconfig:
            data_base = filecoreconfig.read()
    filecoreconfig.close()

    with open('oai_db.sql', 'w') as filecoreconfig:
        filecoreconfig.write(data_base)


def gnb_config():

    if os.path.isfile('Templete/ConfigureCore_gNB_template_fast.conf'):
        with open('Templete/ConfigureCore_gNB_template_fast.conf', 'r') as filegnbconfig:
            gnb_data = filegnbconfig.read()
    filegnbconfig.close()

    new_nicgnb = '"'+nicgnb+'";'
    gnb_data = gnb_data.replace('nicgnb', new_nicgnb)

    with open('configure.conf', 'w') as filegnbconfig:
        filegnbconfig.write(gnb_data)


def oaiue_config():
    if os.path.isfile('Templete/ue_fast_templete.conf'):
        with open('Templete/ue_fast_templete.conf', 'r') as fileoaiconfig:
            oaiue_data = fileoaiconfig.read()
    fileoaiconfig.close()

    with open('ue.conf', 'w') as fileoaiconfig:
        fileoaiconfig.write(oaiue_data)

###########################################
#          All in One - Advance           #
###########################################


def core_config_allinone_ad():  # repetido do sleep core/gNB mas para ser rapido agora fica assim, pode ser para apagar
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipcore, username=usercore, password=passcore)

    dirlocal = subprocess.getoutput('pwd')

    cmd = 'netstat -ie | grep -B1 "' + ipcore + '" | head -n1 | awk ' + '\'' + '{print $1}' + '\'' + ';'
    print(cmd)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command(cmd)
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    data = str(data)
    interfacecore = data[2:len(data) - 5]
    print(interfacecore)

    cmd = 'nmcli device show ' + interfacecore + ' | grep IP4.DNS > dns.txt'
    client.exec_command(cmd)

    sftp = client.open_sftp()
    localpath = dirlocal + '/dns.txt'
    remotepath = '/home/' + usercore + '/dns.txt'
    sftp.get(remotepath, localpath)
    sftp.close()

    if os.path.isfile('dns.txt'):
        with open('dns.txt', 'r') as file:
            dns_data = file.read()
    file.close()

    dns1 = str(dns_data[40:53])
    dns2 = str(dns_data[94:len(dns_data)])

    print(dns1)
    print(dns2)

    if os.path.isfile('Templete/docker-compose-basic-nrf_templete_advance.yaml'):
        with open('Templete/docker-compose-basic-nrf_templete_advance.yaml', 'r') as filecoreconfigallinone:
            docker_compose = filecoreconfigallinone.read()
    filecoreconfigallinone.close()
    ##############################################################
    docker_compose = docker_compose.replace('dns1', dns1)
    docker_compose = docker_compose.replace('dns2', dns2)
    docker_compose = docker_compose.replace('mcc', str(mccgnb))
    docker_compose = docker_compose.replace('mnc', str(mncgnb))
    # docker_compose = docker_compose.replace('sub_mnc03', '0'+str(mccgnb))
    ##############################################################
    with open('docker-compose-basic-nrf.yaml', 'w') as filecoreconfigallinone:
        filecoreconfigallinone.write(docker_compose)

    if os.path.isfile('Templete/oai_dbCore_gNB_advance.sql'):  # Atenção alterar para uma dBase propria
        with open('Templete/oai_dbCore_gNB_advance.sql', 'r') as filecoreconfigallinone:
            data_base = file.read()
    file.close()

    data_base = data_base.replace('sub_imsi', str(imsi))
    data_base = data_base.replace('sub_mccmnc', str(mccgnb) + str(mncgnb))
    data_base = data_base.replace('sub_sst', str(sstgnb))
    data_base = data_base.replace('sub_sd', str(sdgnb))
    data_base = data_base.replace('sub_dnn', str(dnn))
    data_base = data_base.replace('sub_key', str(key))

    with open('oai_db.sql', 'w') as filecoreconfigallinone:
        filecoreconfigallinone.write(data_base)


def gnb_config_allinone_ad():
    if os.path.isfile('Templete/ConfigureCore_gNB_template_advance.conf'):
        with open('Templete/ConfigureCore_gNB_template_advance.conf', 'r') as filegnbconfigallinone:
            gnb_data = filegnbconfigallinone.read()
    filegnbconfigallinone.close()

    gnb_data = gnb_data.replace('sub_name', str(namegnb))
    gnb_data = gnb_data.replace('sub_mcc', str(mccgnb))
    gnb_data = gnb_data.replace('sub_mnc', str(mncgnb))
    gnb_data = gnb_data.replace('sub_sst', str(sstgnb))
    gnb_data = gnb_data.replace('sub_sd', str(sdgnb))
    gnb_data = gnb_data.replace('sub_band', str(bandgnb))

    with open('configure.conf', 'w') as filegnbconfigallinone:
        filegnbconfigallinone.write(gnb_data)


def oai_config_allinone_ad():
    if os.path.isfile('Templete/ue_advance_templete.conf'):
        with open('Templete/ue_advance_templete.conf', 'r') as fileoaiconfigallinone:
            ue_data = fileoaiconfigallinone.read()
    fileoaiconfigallinone.close()
    ue_data = ue_data.replace('sub_imsi', str(imsi))
    ue_data = ue_data.replace('sub_key', str(key))
    ue_data = ue_data.replace('sub_opc', str(opc))
    ue_data = ue_data.replace('sub_dnn', str(dnn))
    ue_data = ue_data.replace('sub_sst', str(sstgnb))
    ue_data = ue_data.replace('sub_sd', str(sdgnb))

    with open('ue.conf', 'w') as fileoaiconfigallinone:
        fileoaiconfigallinone.write(ue_data)


###########################################
#          Core e gNB - Advance           #
###########################################

def core_config_ad():
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipcore, username=usercore, password=passcore)

    dirlocal = subprocess.getoutput('pwd')

    cmd = 'netstat -ie | grep -B1 "' + ipcore + '" | head -n1 | awk ' + '\'' + '{print $1}' + '\'' + ';'
    print(cmd)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command(cmd)
    ssh_stdin.flush()

    data = ssh_stdout.readlines()
    data = str(data)
    interfacecore = data[2:len(data) - 5]
    print(interfacecore)

    cmd = 'nmcli device show ' + interfacecore + ' | grep IP4.DNS > dns.txt'
    client.exec_command(cmd)

    sftp = client.open_sftp()
    localpath = dirlocal + '/dns.txt'
    remotepath = '/home/' + usercore + '/dns.txt'
    sftp.get(remotepath, localpath)
    sftp.close()

    if os.path.isfile('dns.txt'):
        with open('dns.txt', 'r') as filecoreconfigad:
            dns_data = filecoreconfigad.read()
    filecoreconfigad.close()

    dns1 = str(dns_data[40:53])
    dns2 = str(dns_data[94:len(dns_data)])

    print(dns1)
    print(dns2)

    if os.path.isfile('Templete/docker-compose-basic-nrf_templete_advance.yaml'):
        with open('Templete/docker-compose-basic-nrf_templete_advance.yaml', 'r') as filecoreconfigad:
            docker_compose = filecoreconfigad.read()
    filecoreconfigad.close()
##############################################################
    docker_compose = docker_compose.replace('dns1', dns1)
    docker_compose = docker_compose.replace('dns2', dns2)
    docker_compose = docker_compose.replace('mcc', str(mccgnb))
    docker_compose = docker_compose.replace('mnc', str(mncgnb))
    # docker_compose = docker_compose.replace('sub_mnc03', '0'+str(mccgnb))
##############################################################
    with open('docker-compose-basic-nrf.yaml', 'w') as filecoreconfigad:
        filecoreconfigad.write(docker_compose)

    if os.path.isfile('Templete/oai_dbCore_gNB_advance.sql'):  # Atenção alterar para uma dBase propria
        with open('Templete/oai_dbCore_gNB_advance.sql', 'r') as filecoreconfigad:
            data_base = filecoreconfigad.read()
    filecoreconfigad.close()

    data_base = data_base.replace('sub_imsi', str(imsi))
    data_base = data_base.replace('sub_mccmnc', str(mccgnb)+str(mncgnb))
    data_base = data_base.replace('sub_sst', str(sstgnb))
    data_base = data_base.replace('sub_sd', str(sdgnb))
    data_base = data_base.replace('sub_dnn', str(dnn))
    data_base = data_base.replace('sub_key', str(key))
    data_base = data_base.replace('sub_opc', str(opc))

    with open('oai_db.sql', 'w') as filecoreconfigad:
        filecoreconfigad.write(data_base)


def gnb_config_ad():
    if os.path.isfile('Templete/ConfigureCore_gNB_template_advance.conf'):
        with open('Templete/ConfigureCore_gNB_template_advance.conf', 'r') as filegnbconfigad:
            gnb_data = filegnbconfigad.read()
    filegnbconfigad.close()

    gnb_data = gnb_data.replace('sub_name', str(namegnb))
    gnb_data = gnb_data.replace('sub_id', str(idgnb))
    new_nicgnb = '"'+nicgnb+'";'
    gnb_data = gnb_data.replace('nicgnb', new_nicgnb)
    gnb_data = gnb_data.replace('sub_mcc', str(mccgnb))
    gnb_data = gnb_data.replace('sub_mnc', str(mncgnb))
    gnb_data = gnb_data.replace('sub_sst', str(sstgnb))
    gnb_data = gnb_data.replace('sub_sd', str(sdgnb))
    gnb_data = gnb_data.replace('sub_band', str(bandgnb))

    with open('configure.conf', 'w') as filegnbconfigad:
        filegnbconfigad.write(gnb_data)


def oaiue_config_ad():
    if os.path.isfile('Templete/ue_advance_templete.conf'):
        with open('Templete/ue_advance_templete.conf', 'r') as fileoaiconfigad:
            ue_data = fileoaiconfigad.read()
    fileoaiconfigad.close()
    ue_data = ue_data.replace('sub_imsi', str(imsi))
    ue_data = ue_data.replace('sub_key', str(key))
    ue_data = ue_data.replace('sub_opc', str(opc))
    ue_data = ue_data.replace('sub_dnn', str(dnn))
    ue_data = ue_data.replace('sub_sst', str(sstgnb))
    ue_data = ue_data.replace('sub_sd', str(sdgnb))

    with open('ue.conf', 'w') as fileoaiconfigad:
        fileoaiconfigad.write(ue_data)

#################################################
#                    Cartão                     #
#################################################


def program_card(adm, imsi, key, opc, spn):
    state = True
    print('Começou a programar')
    cmd = 'echo 67f7fb87 | sudo -S uicc-v2.6/./program_uicc --adm '+adm+' --imsi '+imsi+' --isdn 00000001 --acc 0001 --key '+key+' --opc '+opc+' -spn "'+spn+'" --authenticate --noreadafter'
    print(cmd)
    info = subprocess.getoutput(cmd)
#################################################
#                    Testes                     #
#################################################


def run_gnb_tests():

    if os.path.isfile('Templete/ConfiguregNB_template.conf'):
        with open('Templete/ConfiguregNB_template.conf', 'r') as filerungnb:
            gnb_data = filerungnb.read()
    filerungnb.close()

    gnb_data = gnb_data.replace('sub_gnbname', str(gnbname))
    gnb_data = gnb_data.replace('sub_gnbid', str(gnbid))
    gnb_data = gnb_data.replace('sub_mcc', str(mccgnb))
    gnb_data = gnb_data.replace('sub_mnc', str(mncgnb))
    gnb_data = gnb_data.replace('sub_sst', str(sstgnb))
    gnb_data = gnb_data.replace('sub_sd', str(sdgnb))
    gnb_data = gnb_data.replace('sub_frequencyssb', str(num_absoluteFrequencySSB(band, scs, bw)))
    gnb_data = gnb_data.replace('sub_band', str(band))
    gnb_data = gnb_data.replace('sub_pointa', str(num_dl_absoluteFrequencyPointA(band, scs, bw)))
    gnb_data = gnb_data.replace('sub_scs', str(idx_scs(scs)))
    gnb_data = gnb_data.replace('sub_prb', str(num_prb(scs, bw)))
    gnb_data = gnb_data.replace('sub_ilb', str(num_initialDLULBWPlocationAndBandwidth(num_prb(scs, bw))))
    gnb_data = gnb_data.replace('sub_amfip', str(amfip))
    gnb_data = gnb_data.replace('nameamf', str(gnbnameamf))
    gnb_data = gnb_data.replace('ipamf', str(gnbipamf))
    gnb_data = gnb_data.replace('namengu', str(gnbnamegnu))
    gnb_data = gnb_data.replace('ipngu', str(gnbipgnu))

    with open('configure.conf', 'w') as filerungnb:
        filerungnb.write(gnb_data)

    dirlocal = subprocess.getoutput('pwd')

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipgnb, username=usergnb, password=passgnb)

    sftp = client.open_sftp()
    localpath = dirlocal + '/configure.conf'
    remotepath = '/'+str(dirgnb)+'/openairinterface5g/targets/PROJECTS/GENERIC-NR-5GC/CONF/configure.conf'
    sftp.put(localpath, remotepath)
    sftp.close()

    time.sleep(5)  # dar tempo para atualizar

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('cd /'+str(dirgnb)+'/openairinterface5g/cmake_targets/ran_build/build; sudo -S ./nr-softmodem -O /'+str(dirgnb)+'/openairinterface5g/targets/PROJECTS/GENERIC-NR-5GC/CONF/configure.conf --sa --continuous-tx -E > logs_gNB.txt;')
    ssh_stdin.write(passgnb + '\n')
    ssh_stdin.flush()

    data_gnb = ssh_stdout.readlines()  # resolveu com isto
    ssh_stdin.close()


def run_ue_tests():
    time.sleep(30)
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('cd /'+str(dirue)+'/openairinterface5g/cmake_targets/ran_build/build; sudo -S ./nr-uesoftmodem -r 106 --numerology 1 --band 78 -C 3619200000 --ue-fo-compensation --sa --nokrnmod -O /'+str(dirue)+'/openairinterface5g/targets/PROJECTS/GENERIC-NR-5GC/CONF/ue.conf -E > logs_oaiue.txt;')
    ssh_stdin.write(passoaiue + '\n')
    ssh_stdin.flush()

    data_oaiue = ssh_stdout.readlines()  # resolveu com isto
    ssh_stdin.close()


def coreon_test():
    print('Teste Core on')
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipgnb, username=usergnb, password=passgnb)
    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('ping 192.168.70.131 -c 10')
    ssh_stdin.flush()
    data = ssh_stdout.readlines()

    if os.path.isfile('coreon_result.txt'):
        os.remove('coreon_result.txt')  # apagar o antigo
    for linecoreon in data:
        with open('coreon_result.txt', 'a') as filecoreon:
            filecoreon.write(str(linecoreon))
    ssh_stdin.close()


def ueon_test():
    time.sleep(50)

    print('Teste UE')
    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)
    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('ifconfig oaitun_ue1')
    ssh_stdin.flush()
    data = ssh_stdout.readlines()

    if os.path.isfile('ueon_result.txt'):
        os.remove('ueon_result.txt')  # apagar o antigo
    for lineueon in data:
        with open('ueon_result.txt', 'a') as fileueon:
            fileueon.write(str(lineueon))
    ssh_stdin.close()

    print(data)


def ueconnect_test():
    time.sleep(50)

    print('Teste UE Connect')

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)
    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('ping -I oaitun_ue1 8.8.8.8 -c 10 > ueconnect_result.txt')
    ssh_stdin.flush()
    data = ssh_stdout.readlines()

    ssh_stdin.close()

    dirlocal = subprocess.getoutput('pwd')

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)

    sftp = client.open_sftp()
    localpath = dirlocal + '/ueconnect_result.txt'
    remotepath = '/home/' + useroaiue + '/ueconnect_result.txt'
    sftp.get(remotepath, localpath)
    sftp.close()


def iperf_ue():
    time.sleep(69)
    controlo = False
    with open('ip_test.txt','r') as file:
        ip = file.read()
    ips = open('ips.txt')
    lines = ips.readlines()
    for line in lines:
        if line in ip:
            command = line
            controlo = True
            break
    if os.path.isfile('Templete/Iperf_Test_Template.py'):
        with open('Templete/Iperf_Test_Template.py', 'r') as file:
            ue_data = file.read()
    file.close()

    if controlo:

        print(command)

        command = command.strip('\n')

        cmd2 = 'result_iperf = subprocess.getoutput('+'\'' + 'iperf -s -u -i 1 -t 20 -B ' + command + '\''+')'


        cmd1 = 'os.system('+'\''+'echo '+str(passoaiue)+' | sudo -S route add default gw '+command+'\''+')'

        print(cmd1)
        print(cmd2)

        ue_data = ue_data.replace('sub_cmd2', cmd2)
        ue_data = ue_data.replace('sub_cmd1', cmd1)

    with open('Iperf_Test.py', 'w') as file:
        file.write(ue_data)

    dirlocal = subprocess.getoutput('pwd')

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)

    sftp = client.open_sftp()
    localpath = dirlocal + '/Iperf_Test.py'
    remotepath = '/home/' + useroaiue + '/Iperf_Test.py'
    sftp.put(localpath, remotepath)
    sftp.close()

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('python3 Iperf_Test.py > RESULTADO.txt')
    ssh_stdin.flush()
    data = ssh_stdout.readlines()

    print(data)

    with open('stop_ue.txt', 'a') as file:
        file.write(str('line'))


def iperf_core():
    with open('ip_test.txt', 'r') as file:
        ip = file.read()
    ips = open('ips.txt')
    lines = ips.readlines()
    for line in lines:
        if line in ip:
            break

    cmd = 'docker exec oai-ext-dn iperf -u -t 20 -i 1 -fk -B 192.168.70.135 -b 20M -c ' + line

    with open('stop_core.txt', 'a') as file:
        file.write(str('line'))

    time.sleep(20)

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipcore, username=usercore, password=passcore)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command(cmd)
    ssh_stdin.flush()
    data = ssh_stdout.readlines()

    print(data)


def iperf_test():
    time.sleep(50)

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)

    ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command('ifconfig oaitun_ue1')
    ssh_stdin.flush()
    data = ssh_stdout.readlines()
    if os.path.isfile('ip_test.txt'):
        os.remove('ip_test.txt')  # apagar o antigo
    for line in data:
        with open('ip_test.txt', 'a') as file:
            file.write(str(line))
    ssh_stdin.close()

    time.sleep(1)

    print('ENtrouuuuu')

#####################################
    if os.path.isfile('ip_test.txt'):
        with open('ip_test.txt', 'r') as file:
            ip = file.read()
        ips = open('ips.txt')
        lines = ips.readlines()
        for line in lines:
            if line in ip:
                break

        cmd = 'docker exec oai-ext-dn iperf -u -t 20 -i 1 -fk -B 192.168.70.135 -b 20M -c ' + line

        time.sleep(20) # 60

        client = paramiko.SSHClient()
        client.load_system_host_keys()
        client.connect(ipcore, username=usercore, password=passcore)

        ssh_stdin, ssh_stdout, ssh_stderr = client.exec_command(cmd)
        ssh_stdin.flush()
        data = ssh_stdout.readlines()


    time.sleep(25)  # 60

    threading.Thread(target=stop_gnb).start()

    time.sleep(20)

    dirlocal = subprocess.getoutput('pwd')

    client = paramiko.SSHClient()
    client.load_system_host_keys()
    client.connect(ipoaiue, username=useroaiue, password=passoaiue)

    sftp = client.open_sftp()
    localpath = dirlocal + '/result_iperf.txt'
    remotepath = '/home/' + useroaiue + '/result_iperf.txt'
    sftp.get(remotepath, localpath)
    sftp.close()

    with open('stop.txt', 'a') as file:
        file.write(str('line'))

#############################################################
#                  Relatorios dso testes                    #
#############################################################


def relatorio_coreconnect_certo(diretorio):
    data_hora = subprocess.getoutput('date')

    doc = Document('Templete/RelatorioCoreConection.docx')
    for p in doc.paragraphs:  # data
        if 'Data:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Data:' in inline[i].text:
                    text = inline[i].text.replace('Data:', 'Data: ' + str(data_hora))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip core
        if 'Endereço de ip da maquina que corre o core (NIC1):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o core (NIC1):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o core (NIC1):',
                                                  'Endereço de ip da maquina que corre o core (NIC1): ' + str(ipcore))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip do conteiner do amf:' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'Endereço de ip do conteiner do amf:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip do conteiner do amf:',
                                                  'Endereço de ip do conteiner do amf: ' + str('192.168.70.132'))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip da maquina de teste (NIC2):' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'Endereço de ip da maquina de teste (NIC2):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina de teste (NIC2):',
                                                  'Endereço de ip da maquina de teste (NIC2): ' + str(ipgnb))
                    inline[i].text = text
    doc.save('dest1.docx')

    data = subprocess.getoutput('date +%d%B')
    hora = subprocess.getoutput('date +%R')
    hora = hora.replace(':', '_')
    nome = data + '_' + hora

    subprocess.getoutput('mkdir ~/Desktop/Logs/' + str(nome))

    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'dest1.docx'])

    subprocess.getoutput('mv dest1.pdf Relatorio_' + str(nome) + '.pdf')

    cmd = 'mv Relatorio_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
    os.system(cmd)


def relatorio_coreconnect_erro(diretorio):
    data_hora = subprocess.getoutput('date')
    print(data_hora)

    doc = Document('Templete/RelatorioCoreConection_erro.docx')
    for p in doc.paragraphs:  # data
        if 'Data:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Data:' in inline[i].text:
                    text = inline[i].text.replace('Data:', 'Data: ' + str(data_hora))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip core
        if 'Endereço de ip da maquina que corre o core (NIC1):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o core (NIC1):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o core (NIC1):',
                                                  'Endereço de ip da maquina que corre o core (NIC1): ' + str(ipcore))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip do conteiner do amf:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip do conteiner do amf:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip do conteiner do amf:',
                                                  'Endereço de ip do conteiner do amf: ' + str('192.168.70.132'))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip da maquina de teste (NIC2):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina de teste (NIC2):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina de teste (NIC2):',
                                                  'Endereço de ip da maquina de teste (NIC2): ' + str(ipgnb))
                    inline[i].text = text
    doc.save('dest1.docx')

    data = subprocess.getoutput('date +%d%B')
    hora = subprocess.getoutput('date +%R')
    hora = hora.replace(':', '_')
    nome = data + '_' + hora

    subprocess.getoutput('mkdir ~/Desktop/Logs/' + str(nome))

    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'dest1.docx'])

    subprocess.getoutput('mv dest1.pdf Relatorio_' + str(nome) + '.pdf')

    cmd = 'mv Relatorio_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
    os.system(cmd)


def relatorio_ueconnect_certo(diretorio, band, scs, bw, line):
    data_hora = subprocess.getoutput('date')
    print(data_hora)

    doc = Document('Templete/RelatorioUEConection.docx')
    for p in doc.paragraphs:  # data
        if 'Data:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Data:' in inline[i].text:
                    text = inline[i].text.replace('Data:', 'Data: ' + str(data_hora))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip core
        if 'Endereço de ip da maquina que corre o core (NIC1):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o core (NIC1):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o core (NIC1):',
                                                  'Endereço de ip da maquina que corre o core (NIC1): ' + str(ipcore))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip do conteiner do amf:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip do conteiner do amf:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip do conteiner do amf:',
                                                  'Endereço de ip do conteiner do amf: ' + str('192.168.70.132'))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip gnb
        if 'Endereço de ip da maquina que corre o gNB (NIC2):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o gNB (NIC2):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o gNB (NIC2):',
                                                  'Endereço de ip da maquina que corre o gNB (NIC2): ' + str(ipgnb))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # band
        if 'Band:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Band:' in inline[i].text:
                    text = inline[i].text.replace('Band:', 'Band: ' + str(band))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # scs
        if 'Subcarrier Spacing (SCS):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Subcarrier Spacing (SCS):' in inline[i].text:
                    text = inline[i].text.replace('Subcarrier Spacing (SCS):', 'Subcarrier Spacing (SCS): ' + str(scs))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # bw
        if 'BandWidth (BW):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'BandWidth (BW):' in inline[i].text:
                    text = inline[i].text.replace('BandWidth (BW):', 'BandWidth (BW): ' + str(bw))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Endereço de ip da maquina que corre o UE (NIC3):' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o UE (NIC3):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o UE (NIC3):',
                                                  'Endereço de ip da maquina que corre o UE (NIC3): ' + str(ipoaiue))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Endereço de ip atribuído ao UE pelo Core:' in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if 'Endereço de ip atribuído ao UE pelo Core:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip atribuído ao UE pelo Core:',
                                                  'Endereço de ip atribuído ao UE pelo Core: ' + str(line))
                    inline[i].text = text

    doc.save('dest1.docx')

    data = subprocess.getoutput('date +%d%B')
    hora = subprocess.getoutput('date +%R')
    hora = hora.replace(':', '_')
    nome = data + '_' + hora

    subprocess.getoutput('mkdir ~/Desktop/Logs/' + str(nome))

    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'dest1.docx'])

    subprocess.getoutput('mv dest1.pdf Relatorio_' + str(nome) + '.pdf')

    cmd = 'mv Relatorio_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
    os.system(cmd)

    if relgnb:
        subprocess.getoutput('mv Logs_gNB.pdf Logs_gNB_' + str(nome) + '.pdf')
        cmd = 'mv Logs_gNB_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)

    if relue:
        subprocess.getoutput('mv Logs_oaiue.pdf Logs_oaiue_' + str(nome) + '.pdf')
        cmd = 'mv Logs_oaiue_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)


def relatorio_ueconnect_erro(diretorio, band, scs, bw):
    data_hora = subprocess.getoutput('date')
    print(data_hora)

    doc = Document('Templete/RelatorioUEConection_erro.docx')
    for p in doc.paragraphs:  # data
        if 'Data:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Data:' in inline[i].text:
                    text = inline[i].text.replace('Data:', 'Data: ' + str(data_hora))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip core
        if 'Endereço de ip da maquina que corre o core (NIC1):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o core (NIC1):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o core (NIC1):',
                                                  'Endereço de ip da maquina que corre o core (NIC1): ' + str(ipcore))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip do conteiner do amf:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip do conteiner do amf:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip do conteiner do amf:',
                                                  'Endereço de ip do conteiner do amf: ' + str('192.168.70.132'))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip gnb
        if 'Endereço de ip da maquina que corre o gNB (NIC2):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o gNB (NIC2):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o gNB (NIC2):',
                                                  'Endereço de ip da maquina que corre o gNB (NIC2): ' + str(ipgnb))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # band
        if 'Band:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Band:' in inline[i].text:
                    text = inline[i].text.replace('Band:', 'Band: ' + str(band))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # scs
        if 'Subcarrier Spacing (SCS):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Subcarrier Spacing (SCS):' in inline[i].text:
                    text = inline[i].text.replace('Subcarrier Spacing (SCS):',
                                                  'Subcarrier Spacing (SCS): ' + str(scs))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # bw
        if 'BandWidth (BW):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'BandWidth (BW):' in inline[i].text:
                    text = inline[i].text.replace('BandWidth (BW):', 'BandWidth (BW): ' + str(bw))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Endereço de ip da maquina que corre o UE (NIC3):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o UE (NIC3):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o UE (NIC3):',
                                                  'Endereço de ip da maquina que corre o UE (NIC3): ' + str(ipoaiue))
                    inline[i].text = text
    doc.save('dest1.docx')

    data = subprocess.getoutput('date +%d%B')
    hora = subprocess.getoutput('date +%R')
    hora = hora.replace(':', '_')
    nome = data + '_' + hora

    subprocess.getoutput('mkdir ~/Desktop/Logs/' + str(nome))

    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'dest1.docx'])

    subprocess.getoutput('mv dest1.pdf Relatorio_' + str(nome) + '.pdf')

    cmd = 'mv Relatorio_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
    os.system(cmd)

    if relgnb:
        subprocess.getoutput('mv Logs_gNB.pdf Logs_gNB_' + str(nome) + '.pdf')
        cmd = 'mv Logs_gNB_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)

    if relue:
        subprocess.getoutput('mv Logs_oaiue.pdf Logs_oaiue_' + str(nome) + '.pdf')
        cmd = 'mv Logs_oaiue_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)


def relatorio_uetoext_certo(diretorio, band, scs, bw, line):
    data_hora = subprocess.getoutput('date')

    doc = Document('Templete/RelatorioConectionOutside.docx')
    for p in doc.paragraphs:  # data
        if 'Data:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Data:' in inline[i].text:
                    text = inline[i].text.replace('Data:', 'Data: ' + str(data_hora))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip core
        if 'Endereço de ip da maquina que corre o core (NIC1):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o core (NIC1):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o core (NIC1):',
                                                  'Endereço de ip da maquina que corre o core (NIC1): ' + str(ipcore))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip do conteiner do amf:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip do conteiner do amf:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip do conteiner do amf:',
                                                  'Endereço de ip do conteiner do amf: ' + str('192.168.70.132'))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip gnb
        if 'Endereço de ip da maquina que corre o gNB (NIC2):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o gNB (NIC2):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o gNB (NIC2):',
                                                  'Endereço de ip da maquina que corre o gNB (NIC2): ' + str(ipgnb))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # band
        if 'Band:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Band:' in inline[i].text:
                    text = inline[i].text.replace('Band:', 'Band: ' + str(band))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # scs
        if 'Subcarrier Spacing (SCS):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Subcarrier Spacing (SCS):' in inline[i].text:
                    text = inline[i].text.replace('Subcarrier Spacing (SCS):',
                                                  'Subcarrier Spacing (SCS): ' + str(scs))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # bw
        if 'BandWidth (BW):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'BandWidth (BW):' in inline[i].text:
                    text = inline[i].text.replace('BandWidth (BW):', 'BandWidth (BW): ' + str(bw))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Endereço de ip da maquina que corre o UE (NIC3):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o UE (NIC3):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o UE (NIC3):',
                                                  'Endereço de ip da maquina que corre o UE (NIC3): ' + str(ipoaiue))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Endereço de ip atribuído ao UE pelo Core:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip atribuído ao UE pelo Core:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip atribuído ao UE pelo Core:',
                                                  'Endereço de ip atribuído ao UE pelo Core: ' + str(line))
                    inline[i].text = text

    doc.save('dest1.docx')

    data = subprocess.getoutput('date +%d%B')
    hora = subprocess.getoutput('date +%R')
    hora = hora.replace(':', '_')
    nome = data + '_' + hora

    subprocess.getoutput('mkdir ~/Desktop/Logs/' + str(nome))

    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'dest1.docx'])

    subprocess.getoutput('mv dest1.pdf Relatorio_' + str(nome) + '.pdf')

    cmd = 'mv Relatorio_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
    os.system(cmd)

    if relgnb:
        subprocess.getoutput('mv Logs_gNB.pdf Logs_gNB_' + str(nome) + '.pdf')
        cmd = 'mv Logs_gNB_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)

    if relue:
        subprocess.getoutput('mv Logs_oaiue.pdf Logs_oaiue_' + str(nome) + '.pdf')
        cmd = 'mv Logs_oaiue_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)


def relatorio_uetoext_erro(diretorio, band, scs, bw):
    data_hora = subprocess.getoutput('date')
    print(data_hora)

    doc = Document('Templete/RelatorioConectionOutside_erro.docx')
    for p in doc.paragraphs:  # data
        if 'Data:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Data:' in inline[i].text:
                    text = inline[i].text.replace('Data:', 'Data: ' + str(data_hora))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip core
        if 'Endereço de ip da maquina que corre o core (NIC1):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o core (NIC1):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o core (NIC1):',
                                                  'Endereço de ip da maquina que corre o core (NIC1): ' + str(ipcore))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip do conteiner do amf:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip do conteiner do amf:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip do conteiner do amf:',
                                                  'Endereço de ip do conteiner do amf: ' + str('192.168.70.132'))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip gnb
        if 'Endereço de ip da maquina que corre o gNB (NIC2):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o gNB (NIC2):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o gNB (NIC2):',
                                                  'Endereço de ip da maquina que corre o gNB (NIC2): ' + str(ipgnb))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # band
        if 'Band:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Band:' in inline[i].text:
                    text = inline[i].text.replace('Band:', 'Band: ' + str(band))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # scs
        if 'Subcarrier Spacing (SCS):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Subcarrier Spacing (SCS):' in inline[i].text:
                    text = inline[i].text.replace('Subcarrier Spacing (SCS):',
                                                  'Subcarrier Spacing (SCS): ' + str(scs))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # bw
        if 'BandWidth (BW):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'BandWidth (BW):' in inline[i].text:
                    text = inline[i].text.replace('BandWidth (BW):', 'BandWidth (BW): ' + str(bw))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Endereço de ip da maquina que corre o UE (NIC3):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o UE (NIC3):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o UE (NIC3):',
                                                  'Endereço de ip da maquina que corre o UE (NIC3): ' + str(ipoaiue))
                    inline[i].text = text
    doc.save('dest1.docx')

    data = subprocess.getoutput('date +%d%B')
    hora = subprocess.getoutput('date +%R')
    hora = hora.replace(':', '_')
    nome = data + '_' + hora

    subprocess.getoutput('mkdir ~/Desktop/Logs/' + str(nome))

    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'dest1.docx'])

    subprocess.getoutput('mv dest1.pdf Relatorio_' + str(nome) + '.pdf')

    cmd = 'mv Relatorio_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
    os.system(cmd)

    if relgnb:
        subprocess.getoutput('mv Logs_gNB.pdf Logs_gNB_' + str(nome) + '.pdf')
        cmd = 'mv Logs_gNB_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)

    if relue:
        subprocess.getoutput('mv Logs_oaiue.pdf Logs_oaiue_' + str(nome) + '.pdf')
        cmd = 'mv Logs_oaiue_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)


def relatorio_iperftest_certo(diretorio, band, scs, bw, line, tipo, quantidade, perdidos, total, percentagem):
    data_hora = subprocess.getoutput('date')
    doc = Document('Templete/RelatorioIPerf.docx')
    for p in doc.paragraphs:  # data
        if 'Data:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Data:' in inline[i].text:
                    text = inline[i].text.replace('Data:', 'Data: ' + str(data_hora))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip core
        if 'Endereço de ip da maquina que corre o core (NIC1):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o core (NIC1):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o core (NIC1):',
                                                  'Endereço de ip da maquina que corre o core (NIC1): ' + str(ipcore))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip do conteiner do amf:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip do conteiner do amf:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip do conteiner do amf:',
                                                  'Endereço de ip do conteiner do amf: ' + str('192.168.70.132'))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip gnb
        if 'Endereço de ip da maquina que corre o gNB (NIC2):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o gNB (NIC2):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o gNB (NIC2):',
                                                  'Endereço de ip da maquina que corre o gNB (NIC2): ' + str(ipgnb))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # band
        if 'Band:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Band:' in inline[i].text:
                    text = inline[i].text.replace('Band:', 'Band: ' + str(band))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # scs
        if 'Subcarrier Spacing (SCS):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Subcarrier Spacing (SCS):' in inline[i].text:
                    text = inline[i].text.replace('Subcarrier Spacing (SCS):',
                                                  'Subcarrier Spacing (SCS): ' + str(scs))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # bw
        if 'BandWidth (BW):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'BandWidth (BW):' in inline[i].text:
                    text = inline[i].text.replace('BandWidth (BW):', 'BandWidth (BW): ' + str(bw))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Endereço de ip da maquina que corre o UE (NIC3):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o UE (NIC3):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o UE (NIC3):',
                                                  'Endereço de ip da maquina que corre o UE (NIC3): ' + str(ipoaiue))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Endereço de ip atribuído ao UE pelo Core:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip atribuído ao UE pelo Core:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip atribuído ao UE pelo Core:',
                                                  'Endereço de ip atribuído ao UE pelo Core: ' + str(line))
                    inline[i].text = text

    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Tipo de tráfego:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Tipo de tráfego:' in inline[i].text:
                    text = inline[i].text.replace('Tipo de tráfego:',
                                                  'Tipo de tráfego: ' + str(tipo))
                    inline[i].text = text

    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Total de dados transferidos:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Total de dados transferidos:' in inline[i].text:
                    text = inline[i].text.replace('Total de dados transferidos:',
                                                  'Total de dados transferidos: ' + str(quantidade))
                    inline[i].text = text

    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Numero de pacotes perdidos:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Numero de pacotes perdidos:' in inline[i].text:
                    text = inline[i].text.replace('Numero de pacotes perdidos:',
                                                  'Numero de pacotes perdidos: ' + str(perdidos))
                    inline[i].text = text

    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Total de pacotes:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Total de pacotes:' in inline[i].text:
                    text = inline[i].text.replace('Total de pacotes:',
                                                  'Total de pacotes: ' + str(total))
                    inline[i].text = text

    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Percentagem de pacotes perdidos:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Percentagem de pacotes perdidos:' in inline[i].text:
                    text = inline[i].text.replace('Percentagem de pacotes perdidos:',
                                                  'Percentagem de pacotes perdidos: ' + str(percentagem))
                    inline[i].text = text

    doc.save('dest1.docx')

    data = subprocess.getoutput('date +%d%B')
    hora = subprocess.getoutput('date +%R')
    hora = hora.replace(':', '_')
    nome = data + '_' + hora

    subprocess.getoutput('mkdir ~/Desktop/Logs/' + str(nome))

    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'dest1.docx'])

    subprocess.getoutput('mv dest1.pdf Relatorio_' + str(nome) + '.pdf')

    cmd = 'mv Relatorio_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
    os.system(cmd)

    if relgnb:
        subprocess.getoutput('mv Logs_gNB.pdf Logs_gNB_' + str(nome) + '.pdf')
        cmd = 'mv Logs_gNB_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)

    if relue:
        subprocess.getoutput('mv Logs_oaiue.pdf Logs_oaiue_' + str(nome) + '.pdf')
        cmd = 'mv Logs_oaiue_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)


def relotiro_iperftest_erro(diretorio, band, scs, bw):
    data_hora = subprocess.getoutput('date')
    doc = Document('Templete/RelatorioIPerf_erro.docx')
    for p in doc.paragraphs:  # data
        if 'Data:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Data:' in inline[i].text:
                    text = inline[i].text.replace('Data:', 'Data: ' + str(data_hora))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip core
        if 'Endereço de ip da maquina que corre o core (NIC1):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o core (NIC1):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o core (NIC1):',
                                                  'Endereço de ip da maquina que corre o core (NIC1): ' + str(ipcore))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip amf
        if 'Endereço de ip do conteiner do amf:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip do conteiner do amf:' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip do conteiner do amf:',
                                                  'Endereço de ip do conteiner do amf: ' + str('192.168.70.132'))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ip gnb
        if 'Endereço de ip da maquina que corre o gNB (NIC2):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o gNB (NIC2):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o gNB (NIC2):',
                                                  'Endereço de ip da maquina que corre o gNB (NIC2): ' + str(ipgnb))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # band
        if 'Band:' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Band:' in inline[i].text:
                    text = inline[i].text.replace('Band:', 'Band: ' + str(band))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # scs
        if 'Subcarrier Spacing (SCS):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Subcarrier Spacing (SCS):' in inline[i].text:
                    text = inline[i].text.replace('Subcarrier Spacing (SCS):',
                                                  'Subcarrier Spacing (SCS): ' + str(scs))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # bw
        if 'BandWidth (BW):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'BandWidth (BW):' in inline[i].text:
                    text = inline[i].text.replace('BandWidth (BW):', 'BandWidth (BW): ' + str(bw))
                    inline[i].text = text
    doc.save('dest1.docx')

    doc = Document('dest1.docx')
    for p in doc.paragraphs:  # ue ip
        if 'Endereço de ip da maquina que corre o UE (NIC3):' in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if 'Endereço de ip da maquina que corre o UE (NIC3):' in inline[i].text:
                    text = inline[i].text.replace('Endereço de ip da maquina que corre o UE (NIC3):',
                                                  'Endereço de ip da maquina que corre o UE (NIC3): ' + str(ipoaiue))
                    inline[i].text = text
    doc.save('dest1.docx')

    data = subprocess.getoutput('date +%d%B')
    hora = subprocess.getoutput('date +%R')
    hora = hora.replace(':', '_')
    nome = data + '_' + hora

    subprocess.getoutput('mkdir ~/Desktop/Logs/' + str(nome))

    subprocess.check_output(['libreoffice', '--convert-to', 'pdf', 'dest1.docx'])

    subprocess.getoutput('mv dest1.pdf Relatorio_' + str(nome) + '.pdf')

    cmd = 'mv Relatorio_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
    os.system(cmd)

    if relgnb:
        subprocess.getoutput('mv Logs_gNB.pdf Logs_gNB_' + str(nome) + '.pdf')
        cmd = 'mv Logs_gNB_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)

    if relue:
        subprocess.getoutput('mv Logs_oaiue.pdf Logs_oaiue_' + str(nome) + '.pdf')
        cmd = 'mv Logs_oaiue_' + str(nome) + '.pdf' + ' ~/' + str(diretorio) + '/' + str(nome)
        os.system(cmd)


#################################################################################################
#                                  Inicio do Codigo                                             #
#################################################################################################


# plataformapass = getpass()
limpar()
info = ["null", "null", "null", "null", "null", "null", "null", "null", "null", "null", "null", "null", "null",
        "null", "null", "null", "null", "null", "null", "null", "null", "null", "null", "null", "null", "null",
        "null", "null", "null", "null", "null"]
file1 = open('info.txt', 'r')
lines = file1.readlines()

count = 0
for line in lines:
    info[count] = format(line.strip())
    count += 1

############################################################################
show1 = False
show2 = False
show3 = False
show4 = False
show5 = False
show6 = False
show7 = False
show8 = False

oldinfo = False
#################################################
#          Variaveis (installation)             #
#################################################


terminar = False  # arranjar solução melhor
doinstallcore = False
doinstallgnb = False
doinstalloaiue = False
doinstallcoregnb_oaiue = False
doinstallcore_gnb_oaiue = False
doinstallcoregnb_realue = False
doinstallcore_gnb_realue = False
advancedeployment = False

#################################################
#            Variaveis (Network)                #
#################################################
preparar = False
gnb_on = False
ue_on = False

#################################################
#            Variaveis (Testes)                 #
#################################################
coreon = False
ueon = False
ueconnect = False
iperf = False
result_test_coreon = False
result_test_ueon = False
result_test_ueconnect = False
result_test_iperf = False
preparar_ueon = False
preparar_ueconnect = False
preparar_iperf = False
relgnb = False
relue = False

#################################################
#             Variaveis (Iperf)                 #
#################################################
command = 'nada'


janela1, janela2, janela3, janela4, janela5, janela6, janela7, janela8, janela9, janela10, janela11, janela12, \
janela13, janela14, janela15, janela16, janela17, janela18, janela19, janela20, janela21, janela22, janela23, janela24,\
janela25, janela26, janela27, janela28, janela29, janela30, janela31, janela32, janela33, janela34, janela35, janela36,\
janela37, janela38, janela39, janela40, janela41, janela42, janela43, janela44, janela45, janela46, janela47, janela48,\
janela49, janela50, janela51, janela52, janela53, janela54, janela55, janela56, janela57, janela58, janela59,\
janela60 = interface1(), None, None, None, None, None, None, None, None, None, None, None, None, None, None,\
           None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None,\
           None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, \
           None, None, None, None, None, None, None, None, None

while True:
    window, event, values = sg.read_all_windows(timeout=100)

    info[12] = 'gNB' #apagar

    # Installation

    if window == janela1 and event == sg.WIN_CLOSED:
        break

    if window == janela2 and event == sg.WIN_CLOSED:
        break

    if window == janela3 and event == sg.WIN_CLOSED:
        break

    if window == janela4 and event == sg.WIN_CLOSED:
        break

    if window == janela5 and event == sg.WIN_CLOSED:
        break

    if window == janela6 and event == sg.WIN_CLOSED:
        break

    if window == janela7 and event == sg.WIN_CLOSED:
        break

    if window == janela8 and event == sg.WIN_CLOSED:
        break

    if window == janela9 and event == sg.WIN_CLOSED:
        break

    if window == janela10 and event == sg.WIN_CLOSED:
        break

    if window == janela11 and event == sg.WIN_CLOSED:
        break

    if window == janela12 and event == sg.WIN_CLOSED:
        break

    if window == janela13 and event == sg.WIN_CLOSED:
        break

    if window == janela14 and event == sg.WIN_CLOSED:
        break

    if window == janela15 and event == sg.WIN_CLOSED:
        break

    if window == janela16 and event == sg.WIN_CLOSED:
        break

    if window == janela17 and event == sg.WIN_CLOSED:
        break

    if window == janela18 and event == sg.WIN_CLOSED:
        break

    if window == janela19 and event == sg.WIN_CLOSED:
        break

    if window == janela20 and event == sg.WIN_CLOSED:
        break

    if window == janela21 and event == sg.WIN_CLOSED:
        break

    if window == janela22 and event == sg.WIN_CLOSED:
        break

    if window == janela23 and event == sg.WIN_CLOSED:
        break

    if window == janela24 and event == sg.WIN_CLOSED:
        break

    if window == janela25 and event == sg.WIN_CLOSED:
        break

    if window == janela26 and event == sg.WIN_CLOSED:
        break

    if window == janela1 and event == 'Next' and values['installsoft']:
        janela1.hide()
        janela2 = interface2()
        installsoft = values['installsoft']

    if window == janela2 and event == 'Back':
        janela2.hide()
        janela1.un_hide()

    if window == janela2 and event == 'Next' and values['installcomplete']:
        janela2.hide()
        janela3 = interface3()
        installcomplete = values['installcomplete']

    if window == janela2 and event == 'Next' and values['installcore']:
        janela2.hide()
        janela6 = interface6()
        installcore = values['installcore']

    if window == janela2 and event == 'Next' and values['installgnb']:
        janela2.hide()
        janela7 = interface7()
        installcore = values['installgnb']

    if window == janela2 and event == 'Next' and values['installoaiue']:
        janela2.hide()
        janela8 = interface8()
        installcore = values['installoaiue']

    if window == janela4 and event == 'Next' and values['allinoneoaiue']:
        janela4.hide()
        janela9 = interface9()
        allinoneoaiue = values['allinoneoaiue']

    if window == janela4 and event == 'Next' and values['coregnboaiue']:
        janela4.hide()
        janela10 = interface10()
        coregnboaiue = values['coregnboaiue']

    if window == janela5 and event == 'Next' and values['allinonerealue']:
        janela5.hide()
        janela11 = interface11()
        allinoneralue = values['allinonerealue']

    if window == janela5 and event == 'Next' and values['coregnbrealue']:
        janela5.hide()
        janela12 = interface12()
        coregnbrealue = values['coregnbrealue']

    if window == janela3 and event == 'Back':
        janela3.hide()
        janela2.un_hide()

    if window == janela3 and event == 'Next' and values['oaiue']:
        janela3.hide()
        janela4 = interface4()
        oaiue = values['oaiue']

    if window == janela3 and event == 'Next' and values['realue']:
        janela3.hide()
        janela5 = interface5()
        realue = values['realue']

    if window == janela6 and event == 'Next':
        ipcore = str(values['ipcore'])
        usercore = str(values['usercore'])
        passcore = str(values['passcore'])
        info[0] = ipcore
        info[1] = usercore
        info[2] = passcore
        janela6.hide()
        janela13 = interface13()
        corebash()
        threading.Thread(target=do_install_core).start()
        terminar = True
        doinstallcore = True

    if window == janela4 and event == 'Back':
        janela4.hide()
        janela3.un_hide()

    if window == janela5 and event == 'Back':
        janela5.hide()
        janela3.un_hide()

    if window == janela6 and event == 'Back':
        janela6.hide()
        janela2.un_hide()

    if window == janela7 and event == 'Next':
        ipgnb = values['ipgnb']
        usergnb = values['usergnb']
        passgnb = values['passgnb']
        info[3] = ipgnb
        info[4] = usergnb
        info[5] = passgnb
        janela7.hide()
        janela14 = interface14()
        gnbbash()
        threading.Thread(target=do_install_gnb).start()
        terminar = True
        doinstallgnb = True

    if window == janela7 and event == 'Back':
        janela7.hide()
        janela2.un_hide()

    if window == janela8 and event == 'Next':
        ipoaiue = values['ipoaiue']
        useroaiue = values['useroaiue']
        passoaiue = values['passoaiue']
        info[6] = ipoaiue
        info[7] = useroaiue
        info[8] = passoaiue
        janela8.hide()
        janela15 = interface15()
        threading.Thread(target=do_install_oaiue).start()
        terminar = True
        doinstalloaiue = True

    if window == janela8 and event == 'Back':
        janela8.hide()
        janela2.un_hide()

    if window == janela9 and event == 'Next':
        ipcoregnb = values['ipcoregnb']
        usercoregnb = values['usercoregnb']
        passcoregnb = values['passcoregnb']
        ipoaiue = values['ipoaiue']
        useroaiue = values['useroaiue']
        passoaiue = values['passoaiue']
        info[0] = ipcoregnb
        info[1] = usercoregnb
        info[2] = passcoregnb
        info[3] = ipcoregnb
        info[4] = usercoregnb
        info[5] = passcoregnb
        info[6] = ipoaiue
        info[7] = useroaiue
        info[8] = passoaiue
        janela9.hide()
        janela16 = interface16()
        ipcore = ipcoregnb
        usercore = usercoregnb
        passcore = passcoregnb
        ipgnb = ipcoregnb
        usergnb = usercoregnb
        passgnb = passcoregnb
        corebash()
        gnbbash()
        # falta o oaiue
        threading.Thread(target=do_install_coregnb).start()
        threading.Thread(target=do_install_oaiue).start()
        terminar = True
        doinstallcoregnb_oaiue = True

    if window == janela9 and event == 'Back':
        janela9.hide()
        janela4.un_hide()

    if window == janela10 and event == 'Next':
        ipcore = values['ipcore']
        usercore = values['usercore']
        passcore = values['passcore']
        ipgnb = values['ipgnb']
        usergnb = values['usergnb']
        passgnb = values['passgnb']
        ipoaiue = values['ipoaiue']
        useroaiue = values['useroaiue']
        passoaiue = values['passoaiue']
        info[0] = ipcore
        info[1] = usercore
        info[2] = passcore
        info[3] = ipgnb
        info[4] = usergnb
        info[5] = passgnb
        info[6] = ipoaiue
        info[7] = useroaiue
        info[8] = passoaiue
        janela10.hide()
        janela17 = interface17()
        corebash()
        gnbbash()
        # falta o oaiue
        threading.Thread(target=do_install_core).start()
        threading.Thread(target=do_install_gnb).start()
        threading.Thread(target=do_install_oaiue).start()
        terminar = True
        doinstallcore_gnb_oaiue = True

    if window == janela10 and event == 'Back':
        janela10.hide()
        janela4.un_hide()

    if window == janela11 and event == 'Next':
        ipcoregnb = values['ipcoregnb']
        usercoregnb = values['usercoregnb']
        passcoregnb = values['passcoregnb']
        info[0] = ipcoregnb
        info[1] = usercoregnb
        info[2] = passcoregnb
        info[3] = ipcoregnb
        info[4] = usercoregnb
        info[5] = passcoregnb
        janela11.hide()
        janela18 = interface18()
        ipcore = ipcoregnb
        usercore = usercoregnb
        passcore = passcoregnb
        ipgnb = ipcoregnb
        usergnb = usercoregnb
        passgnb = passcoregnb
        corebash()
        gnbbash()
        threading.Thread(target=do_install_coregnb).start()
        terminar = True
        doinstallcoregnb_realue = True

    if window == janela11 and event == 'Back':
        janela11.hide()
        janela5.un_hide()

    if window == janela12 and event == 'Next':
        ipcore = values['ipcore']
        usercore = values['usercore']
        passcore = values['passcore']
        ipgnb = values['ipgnb']
        usergnb = values['usergnb']
        passgnb = values['passgnb']
        info[0] = ipcore
        info[1] = usercore
        info[2] = passcore
        info[3] = ipgnb
        info[4] = usergnb
        info[5] = passgnb
        janela12.hide()
        janela19 = interface19()
        corebash()
        gnbbash()
        threading.Thread(target=do_install_core).start()
        threading.Thread(target=do_install_gnb).start()
        terminar = True
        doinstallcore_gnb_realue = True

    if window == janela12 and event == 'Back':
        janela12.hide()
        janela5.un_hide()

    if terminar:
        if doinstallcore:
            janela13['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)
        if doinstallgnb:
            janela14['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)
        if doinstalloaiue:
            janela15['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)
        if doinstallcoregnb_oaiue:
            janela16['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)
        if doinstallcore_gnb_oaiue:
            janela17['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)
        if doinstallcoregnb_realue:
            janela18['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)
        if doinstallcore_gnb_realue:
            janela19['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)

    if terminar and doinstallcore and os.path.isfile('finalizadocore.txt'):
        janela13.hide()
        janela20 = interface20()
        terminar = False
        doinstallcore = False

    if terminar and doinstallgnb and os.path.isfile('finalizadognb.txt'):
        janela14.hide()
        janela21 = interface21()
        terminar = False
        doinstallgnb = False

    if terminar and doinstalloaiue and os.path.isfile('finalizadooaiue.txt'):
        janela15.hide()
        janela22 = interface22()
        terminar = False
        doinstalloaiue = False

    if terminar and doinstallcoregnb_oaiue and os.path.isfile('finalizadocore.txt') and os.path.isfile('finalizadognb.txt') and os.path.isfile('finalizadooaiue.txt'):
        janela16.hide()
        janela23 = interface23()
        terminar = False
        doinstallcoregnb_oaiue = False

    if terminar and doinstallcore_gnb_oaiue and os.path.isfile('finalizadocore.txt') and os.path.isfile('finalizadognb.txt') and os.path.isfile('finalizadooaiue.txt'):
        janela17.hide()
        janela24 = interface24()
        terminar = False
        doinstallcore_gnb_oaiue = False

    if terminar and doinstallcoregnb_realue and os.path.isfile('finalizadocore.txt') and os.path.isfile('finalizadognb.txt'):
        janela18.hide()
        janela25 = interface25()
        terminar = False
        doinstallcoregnb_realue = False

    if terminar and doinstallcore_gnb_realue and os.path.isfile('finalizadocore.txt') and os.path.isfile('finalizadognb.txt'):
        janela19.hide()
        janela26 = interface26()
        terminar = False
        doinstallcore_gnb_realue = False

    if window == janela20 and event == 'Close':
        janela20.hide()
        janela1.un_hide()

    if window == janela21 and event == 'Close':
        janela21.hide()
        janela1.un_hide()

    if window == janela22 and event == 'Close':
        janela22.hide()
        janela1.un_hide()

    if window == janela23 and event == 'Close':
        janela23.hide()
        janela1.un_hide()

    if window == janela24 and event == 'Close':
        janela24.hide()
        janela1.un_hide()

    if window == janela25 and event == 'Close':
        janela25.hide()
        janela1.un_hide()

    if window == janela26 and event == 'Close':
        janela26.hide()
        janela1.un_hide()

    # Network (basic)

    if window == janela1 and event == 'Next' and values['fastdeployment']:
        janela1.hide()
        janela27 = interface27()
        installsoft = values['fastdeployment']

    if window == janela27 and event == sg.WIN_CLOSED:
        break

    if window == janela28 and event == sg.WIN_CLOSED:
        break

    if window == janela29 and event == sg.WIN_CLOSED:
        break

    if window == janela30 and event == sg.WIN_CLOSED:
        break

    if window == janela31 and event == sg.WIN_CLOSED:
        break

    if window == janela32 and event == sg.WIN_CLOSED:
        break

    if window == janela33 and event == sg.WIN_CLOSED:
        break

    if window == janela34 and event == sg.WIN_CLOSED:
        break

    if window == janela35 and event == sg.WIN_CLOSED:
        break

    if window == janela27 and event == 'Next':
        allinone = values['allinone']
        gnbcore = values['coregnb']
        janela27.hide()
        janela28 = interface28()

    if window == janela28 and event == 'Next':
        oaiue = values['oaiue']
        realue = values['realue']
        janela28.hide()
        if allinone and oaiue:
            janela29 = interface29()
            janela29['last_ipcore'].update(value=info[0])
            janela29['last_usercore'].update(value=info[1])
            #janela29['last_passcore'].update(value=info[2])
            janela29['last_passcore'].update(value='*******')
            janela29['last_dircore'].update(value=info[9])
            janela29['last_ipoaiue'].update(value=info[6])
            janela29['last_useroaiue'].update(value=info[7])
            #janela29['last_passoaiue'].update(value=info[8])
            janela29['last_passoaiue'].update(value='*******')
            janela29['last_diroaiue'].update(value=info[11])
        if allinone and realue:
            janela30 = interface30()
            janela30['last_ipcore'].update(value=info[0])
            janela30['last_usercore'].update(value=info[1])
            #janela30['last_passcore'].update(value=info[2])
            janela30['last_passcore'].update(value='*******')
            janela30['last_dircore'].update(value=info[9])
        if gnbcore and oaiue:
            janela31 = interface31()
            janela31['last_ipcore'].update(value=info[0])
            janela31['last_usercore'].update(value=info[1])
            #janela31['last_passcore'].update(value=info[2])
            janela31['last_passcore'].update(value='********')
            janela31['last_dircore'].update(value=info[9])
            janela31['last_ipgnb'].update(value=info[3])
            janela31['last_usergnb'].update(value=info[4])
            #janela31['last_passgnb'].update(value=info[5])
            janela31['last_passgnb'].update(value='*******')
            janela31['last_dirgnb'].update(value=info[10])
            janela31['last_ipoaiue'].update(value=info[6])
            janela31['last_useroaiue'].update(value=info[7])
            #janela31['last_passoaiue'].update(value=info[8])
            janela31['last_passoaiue'].update(value='*******')
            janela31['last_diroaiue'].update(value=info[11])
        if gnbcore and realue:
            janela32 = interface32()
            janela32['last_ipcore'].update(value=info[0])
            janela32['last_usercore'].update(value=info[1])
            #janela32['last_passcore'].update(value=info[2])
            janela32['last_passcore'].update(value='********')
            janela32['last_dircore'].update(value=info[9])
            janela32['last_ipgnb'].update(value=info[3])
            janela32['last_usergnb'].update(value=info[4])
            #janela32['last_passgnb'].update(value=info[5])
            janela32['last_passgnb'].update(value='********')
            janela32['last_dirgnb'].update(value=info[10])

    if window == janela29 and event == 'Next':
        ipcore = values['ipcoregnb']
        usercore = values['usercoregnb']
        passcore = values['passcoregnb']
        ipgnb = values['ipcoregnb']
        usergnb = values['usercoregnb']
        passgnb = values['passcoregnb']
        ipoaiue = values['ipoaiue']
        useroaiue = values['useroaiue']
        passoaiue = values['passoaiue']
        dircore = values['dircoregnb']
        dirgnb = values['dircoregnb']
        dirue = values['dirue']
        oldinfo = values['oldinfo']

        if not oldinfo:
            info[0] = ipcore
            info[1] = usercore
            info[2] = passcore
            info[3] = ipgnb
            info[4] = usergnb
            info[5] = passgnb
            info[6] = ipoaiue
            info[7] = useroaiue
            info[8] = passoaiue
            info[9] = dircore
            info[10] = dirgnb
            info[11] = dirue
        if oldinfo:
            ipcore = info[0]
            usercore = info[1]
            passcore = info[2]
            ipgnb = info[3]
            usergnb = info[4]
            passgnb = info[5]
            ipoaiue = info[6]
            useroaiue = info[7]
            passoaiue = info[8]
            dircore = info[9]
            dirgnb = info[10]
            dirue = info[11]
            oldinfo = False

        janela29.hide()
        janela34 = interface34()
        threading.Thread(target=core_config).start()
        show1 = True
        janela34['setuptype'].update(value='The Setup is an All-in-One with OAI UE terminal')
        preparar = True
        threading.Thread(target=run_core).start()
        threading.Thread(target=run_gnb).start()

    if window == janela30 and event == 'Next':
        ipcore = values['ipcoregnb']
        usercore = values['usercoregnb']
        passcore = values['passcoregnb']
        ipgnb = values['ipcoregnb']
        usergnb = values['usercoregnb']
        passgnb = values['passcoregnb']
        dircore = values['dircoregnb']
        dirgnb = values['dircoregnb']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[0] = ipcore
            info[1] = usercore
            info[2] = passcore
            info[3] = ipgnb
            info[4] = usergnb
            info[5] = passgnb
            info[9] = dircore
            info[10] = dirgnb
        if oldinfo:
            ipcore = info[0]
            usercore = info[1]
            passcore = info[2]
            ipgnb = info[3]
            usergnb = info[4]
            passgnb = info[5]
            dircore = info[9]
            dirgnb = info[10]
            oldinfo = False
        janela30.hide()
        janela33 = interface33()

    if window == janela33 and event == 'Next':
        janela33.hide()
        if allinone and realue:
            threading.Thread(target=core_config).start()
            janela34 = interface34()
            show2 = True
            janela34['setuptype'].update(value='The Setup is an All-in-One with COTS UE terminal')
            preparar = True
            threading.Thread(target=run_core).start()
            threading.Thread(target=run_gnb).start()
        if gnbcore and realue:
            nic_core = nic_information(ipcore, usercore, passcore)
            nic_gnb = nic_information(ipgnb, usergnb, passgnb)
            print('Programar o catão')
            adm = values['adm']
            imsi = '208990000000001'
            key = 'fec86ba6eb707ed08905757b1bb44b8f'
            opc = 'C42449363BBAD02B66D16BC975D77CC1'
            spn = 'OpenAirInterface'
            program_card(adm, imsi, key, opc, spn)
            janela35 = interface35()

    if window == janela31 and event == 'Next' and not advancedeployment:
        ipcore = values['ipcore']
        usercore = values['usercore']
        passcore = values['passcore']
        ipgnb = values['ipgnb']
        usergnb = values['usergnb']
        passgnb = values['passgnb']
        ipoaiue = values['ipoaiue']
        useroaiue = values['useroaiue']
        passoaiue = values['passoaiue']
        dircore = values['dircore']
        dirgnb = values['dirgnb']
        dirue = values['dirue']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[0] = ipcore
            info[1] = usercore
            info[2] = passcore
            info[3] = ipgnb
            info[4] = usergnb
            info[5] = passgnb
            info[6] = ipoaiue
            info[7] = useroaiue
            info[8] = passoaiue
            info[9] = dircore
            info[10] = dirgnb
            info[11] = dirue
        if oldinfo:
            ipcore = info[0]
            usercore = info[1]
            passcore = info[2]
            ipgnb = info[3]
            usergnb = info[4]
            passgnb = info[5]
            ipoaiue = info[6]
            useroaiue = info[7]
            passoaiue = info[8]
            dircore = info[9]
            dirgnb = info[10]
            dirue = info[11]
            oldinfo = False
        janela31.hide()
        nic_core = nic_information(ipcore, usercore, passcore)
        nic_gnb = nic_information(ipgnb, usergnb, passgnb)
        janela35 = interface35()

    if window == janela35 and event == 'Next':
        niccore = values['niccore']
        nicgnb = values['nicgnb']
        janela35.hide()
        janela35.hide()
        janela34 = interface34()
        if gnbcore and oaiue:
            threading.Thread(target=core_config).start()
            threading.Thread(target=gnb_config).start()
            threading.Thread(target=oaiue_config).start()
            show3 = True
            janela34['setuptype'].update(value='The setup has a Core and gNB on different machines with OAI UE terminal')
            preparar = True
            threading.Thread(target=run_core).start()
            threading.Thread(target=run_gnb).start()
            threading.Thread(target=run_oaiue).start()

        if gnbcore and realue:
            threading.Thread(target=core_config).start()
            threading.Thread(target=gnb_config).start()
            show4 = True
            janela34['setuptype'].update(value='The setup has a Core and gNB on different machines with COTS UE')
            preparar = True
            threading.Thread(target=run_core).start()
            threading.Thread(target=run_gnb).start()

    if window == janela32 and event == 'Next' and not advancedeployment:
        ipcore = values['ipcore']
        usercore = values['usercore']
        passcore = values['passcore']
        dircore = values['dircore']
        ipgnb = values['ipgnb']
        usergnb = values['usergnb']
        passgnb = values['passgnb']
        dirgnb = values['dirgnb']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[0] = ipcore
            info[1] = usercore
            info[2] = passcore
            info[3] = ipgnb
            info[4] = usergnb
            info[5] = passgnb
            info[9] = dircore
            info[10] = dirgnb
        if oldinfo:
            ipcore = info[0]
            usercore = info[1]
            passcore = info[2]
            ipgnb = info[3]
            usergnb = info[4]
            passgnb = info[5]
            dircore = info[9]
            dirgnb = info[10]
            oldinfo = False
        janela32.hide()
        janela33 = interface33()

    if os.path.isfile('feito_gnb.txt') and os.path.isfile('feito_oaiue.txt') and not advancedeployment:
        janela34.hide()
        janela47 = interface47()
        preparar = False
        os.remove('feito_gnb.txt')
        os.remove('feito_oaiue.txt')

    if os.path.isfile('feito_gnb.txt') and realue and not advancedeployment: #
        janela34.hide()
        janela47 = interface47()    #
        preparar = False    #
        os.remove('feito_gnb.txt')  #
        realue = False  #

    if os.path.isfile('feito_gnb.txt') and os.path.isfile('feito_oaiue.txt') and advancedeployment:
        janela46.hide()
        janela47 = interface47()
        preparar = False
        os.remove('feito_gnb.txt')
        os.remove('feito_oaiue.txt')

    if os.path.isfile('feito_gnb.txt') and realue and advancedeployment: #
        janela46.hide() #
        janela47 = interface47()    #
        preparar = False    #
        os.remove('feito_gnb.txt')  #
        realue = False  #

    if window == janela27 and event == 'Back':
        janela27.hide()
        janela1.un_hide()

    if window == janela28 and event == 'Back':
        janela28.hide()
        janela27.un_hide()

    if window == janela29 and event == 'Back':
        janela29.hide()
        janela28.un_hide()

    if window == janela30 and event == 'Back':
        janela30.hide()
        janela28.un_hide()

    if window == janela31 and event == 'Back':
        janela31.hide()
        janela28.un_hide()

    if window == janela32 and event == 'Back':
        janela32.hide()
        janela28.un_hide()

    if window == janela33 and event == 'Back':
        janela33.hide()
        if allinone and realue:
            janela30.un_hide()
        if gnbcore and realue:
            janela32.un_hide()

    if window == janela35 and event == 'Back':
        janela35.hide()
        if gnbcore and oaiue:
            janela31.un_hide()
        if gnbcore and realue:
            janela33.un_hide()

    if preparar:
        if not advancedeployment:
            janela34['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)
        if advancedeployment:
            janela46['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)

    # Network (advance)

    if window == janela36 and event == sg.WIN_CLOSED:
        break

    if window == janela37 and event == sg.WIN_CLOSED:
        break

    if window == janela38 and event == sg.WIN_CLOSED:
        break

    if window == janela39 and event == sg.WIN_CLOSED:
        break

    if window == janela40 and event == sg.WIN_CLOSED:
        break

    if window == janela41 and event == sg.WIN_CLOSED:
        break

    if window == janela42 and event == sg.WIN_CLOSED:
        break

    if window == janela43 and event == sg.WIN_CLOSED:
        break

    if window == janela44 and event == sg.WIN_CLOSED:
        break

    if window == janela45 and event == sg.WIN_CLOSED:
        break

    if window == janela46 and event == sg.WIN_CLOSED:
        break

    if window == janela47 and event == sg.WIN_CLOSED:
        break

    if window == janela1 and event == 'Next' and values['advancedeployment']:
        janela1.hide()
        janela36 = interface36()
        advancedeployment = values['advancedeployment']

    if window == janela36 and event == 'Next':
        janela36.hide()
        janela37 = interface37()
        allinone = values['allinone']
        coregnb = values['coregnb']

    if window == janela37 and event == 'Next':
        oaiue = values['oaiue']
        realue = values['realue']
        janela37.hide()
        if allinone and oaiue:
            janela38 = interface38()
            janela38['last_ipcore'].update(value=info[0])
            janela38['last_usercore'].update(value=info[1])
            #janela38['last_passcore'].update(value=info[2])
            janela38['last_passcore'].update(value='********')
            janela38['last_dircore'].update(value=info[9])
            janela38['last_ipoaiue'].update(value=info[6])
            janela38['last_useroaiue'].update(value=info[7])
            #janela38['last_passoaiue'].update(value=info[8])
            janela38['last_passoaiue'].update(value='********')
            janela38['last_diroaiue'].update(value=info[11])
        if allinone and realue:
            janela39 = interface39()
            janela39['last_ipcore'].update(value=info[0])
            janela39['last_usercore'].update(value=info[1])
            #janela39['last_passcore'].update(value=info[2])
            janela39['last_passcore'].update(value='********')
            janela39['last_dircore'].update(value=info[9])
        if coregnb and oaiue:
            janela40 = interface40()
            janela40['last_ipcore'].update(value=info[0])
            janela40['last_usercore'].update(value=info[1])
            #janela40['last_passcore'].update(value=info[2])
            janela40['last_passcore'].update(value='********')
            janela40['last_dircore'].update(value=info[9])
            janela40['last_ipgnb'].update(value=info[3])
            janela40['last_usergnb'].update(value=info[4])
            #janela40['last_passgnb'].update(value=info[5])
            janela40['last_passgnb'].update(value='********')
            janela40['last_dirgnb'].update(value=info[10])
            janela40['last_ipoaiue'].update(value=info[6])
            janela40['last_useroaiue'].update(value=info[7])
            #janela40['last_passoaiue'].update(value=info[8])
            janela40['last_passoaiue'].update(value='********')
            janela40['last_diroaiue'].update(value=info[11])
        if coregnb and realue:
            janela41 = interface41()
            janela41['last_ipcore'].update(value=info[0])
            janela41['last_usercore'].update(value=info[1])
            #janela41['last_passcore'].update(value=info[2])
            janela41['last_passcore'].update(value='********')
            janela41['last_dircore'].update(value=info[9])
            janela41['last_ipgnb'].update(value=info[3])
            janela41['last_usergnb'].update(value=info[4])
            #janela41['last_passgnb'].update(value=info[5])
            janela41['last_passgnb'].update(value='********')
            janela41['last_dirgnb'].update(value=info[10])

    if window == janela38 and event == 'Next':
        ipcore = values['ipcoregnb']
        usercore = values['usercoregnb']
        passcore = values['passcoregnb']
        ipgnb = values['ipcoregnb']
        usergnb = values['usercoregnb']
        passgnb = values['passcoregnb']
        ipoaiue = values['ipoaiue']
        useroaiue = values['useroaiue']
        passoaiue = values['passoaiue']
        dircore = values['dircoregnb']
        dirgnb = values['dircoregnb']
        dirue = values['dirue']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[0] = ipcore
            info[1] = usercore
            info[2] = passcore
            info[3] = ipgnb
            info[4] = usergnb
            info[5] = passgnb
            info[6] = ipoaiue
            info[7] = useroaiue
            info[8] = passoaiue
            info[9] = dircore
            info[10] = dirgnb
            info[11] = dirue
        if oldinfo:
            ipcore = info[0]
            usercore = info[1]
            passcore = info[2]
            ipgnb = info[3]
            usergnb = info[4]
            passgnb = info[5]
            ipoaiue = info[6]
            useroaiue = info[7]
            passoaiue = info[8]
            dircore = info[9]
            dirgnb = info[10]
            dirue = info[11]
            oldinfo = False
        janela38.hide()
        janela42 = interface42()
        janela42['last_namegnb'].update(value=info[12])
        janela42['last_bandgnb'].update(value=info[13])
        janela42['last_scsgnb'].update(value=info[14])
        janela42['last_bwgnb'].update(value=info[15])
        janela42['last_mccgnb'].update(value=info[16])
        janela42['last_mncgnb'].update(value=info[17])
        janela42['last_sstgnb'].update(value=info[18])
        janela42['last_sdgnb'].update(value=info[19])
        janela42['last_idgnb'].update(value=info[20])
        janela42['last_imsi'].update(value=info[21])
        janela42['last_key'].update(value=info[22])
        janela42['last_opc'].update(value=info[23])
        janela42['last_dnn'].update(value=info[24])

    if window == janela39 and event == 'Next':
        ipcore = values['ipcoregnb']
        usercore = values['usercoregnb']
        passcore = values['passcoregnb']
        ipgnb = values['ipcoregnb']
        usergnb = values['usercoregnb']
        passgnb = values['passcoregnb']
        dircore = values['dircoregnb']
        dirgnb = values['dircoregnb']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[0] = ipcore
            info[1] = usercore
            info[2] = passcore
            info[3] = ipgnb
            info[4] = usergnb
            info[5] = passgnb
            info[9] = dircore
            info[10] = dirgnb
        if oldinfo:
            ipcore = info[0]
            usercore = info[1]
            passcore = info[2]
            ipgnb = info[3]
            usergnb = info[4]
            passgnb = info[5]
            dircore = info[9]
            dirgnb = info[10]
            oldinfo = False
        janela39.hide()
        janela43 = interface43()
        janela43['last_namegnb'].update(value=info[12])
        janela43['last_bandgnb'].update(value=info[13])
        janela43['last_scsgnb'].update(value=info[14])
        janela43['last_bwgnb'].update(value=info[15])
        janela43['last_mccgnb'].update(value=info[16])
        janela43['last_mncgnb'].update(value=info[17])
        janela43['last_sstgnb'].update(value=info[18])
        janela43['last_sdgnb'].update(value=info[19])
        janela43['last_idgnb'].update(value=info[20])

    if window == janela40 and event == 'Next':
        ipcore = values['ipcore']
        usercore = values['usercore']
        passcore = values['passcore']
        ipgnb = values['ipgnb']
        usergnb = values['usergnb']
        passgnb = values['passgnb']
        ipoaiue = values['ipoaiue']
        useroaiue = values['useroaiue']
        passoaiue = values['passoaiue']
        dircore = values['dircore']
        dirgnb = values['dirgnb']
        dirue = values['dirue']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[0] = ipcore
            info[1] = usercore
            info[2] = passcore
            info[3] = ipgnb
            info[4] = usergnb
            info[5] = passgnb
            info[6] = ipoaiue
            info[7] = useroaiue
            info[8] = passoaiue
            info[9] = dircore
            info[10] = dirgnb
            info[11] = dirue
        if oldinfo:
            ipcore = info[0]
            usercore = info[1]
            passcore = info[2]
            ipgnb = info[3]
            usergnb = info[4]
            passgnb = info[5]
            ipoaiue = info[6]
            useroaiue = info[7]
            passoaiue = info[8]
            dircore = info[9]
            dirgnb = info[10]
            dirue = info[11]
            oldinfo = False
        janela40.hide()
        janela42 = interface42()
        janela42['last_namegnb'].update(value=info[12])
        janela42['last_bandgnb'].update(value=info[13])
        janela42['last_scsgnb'].update(value=info[14])
        janela42['last_bwgnb'].update(value=info[15])
        janela42['last_mccgnb'].update(value=info[16])
        janela42['last_mncgnb'].update(value=info[17])
        janela42['last_sstgnb'].update(value=info[18])
        janela42['last_sdgnb'].update(value=info[19])
        janela42['last_idgnb'].update(value=info[20])
        janela42['last_imsi'].update(value=info[21])
        janela42['last_key'].update(value=info[22])
        janela42['last_opc'].update(value=info[23])
        janela42['last_dnn'].update(value=info[24])

    if window == janela41 and event == 'Next':
        ipcore = values['ipcore']
        usercore = values['usercore']
        passcore = values['passcore']
        ipgnb = values['ipgnb']
        usergnb = values['usergnb']
        passgnb = values['passgnb']
        dircore = values['dircore']
        dirgnb = values['dirgnb']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[0] = ipcore
            info[1] = usercore
            info[2] = passcore
            info[3] = ipgnb
            info[4] = usergnb
            info[5] = passgnb
            info[9] = dircore
            info[10] = dirgnb
        if oldinfo:
            ipcore = info[0]
            usercore = info[1]
            passcore = info[2]
            ipgnb = info[3]
            usergnb = info[4]
            passgnb = info[5]
            dircore = info[9]
            dirgnb = info[10]
            oldinfo = False
        janela41.hide()
        janela43 = interface43()
        janela43['last_namegnb'].update(value=info[12])
        janela43['last_bandgnb'].update(value=info[13])
        janela43['last_scsgnb'].update(value=info[14])
        janela43['last_bwgnb'].update(value=info[15])
        janela43['last_mccgnb'].update(value=info[16])
        janela43['last_mncgnb'].update(value=info[17])
        janela43['last_sstgnb'].update(value=info[18])
        janela43['last_sdgnb'].update(value=info[19])
        janela43['last_idgnb'].update(value=info[20])

    if window == janela42 and event == 'Next':
        namegnb = values['namegnb']
        bandgnb = values['bandgnb']
        scsgnb = values['scsgnb']
        bwgnb = values['bwgnb']
        mccgnb = values['mccgnb']
        mncgnb = values['mncgnb']
        sstgnb = values['sstgnb']
        sdgnb = values['sdgnb']
        idgnb = values['idgnb']
        imsi = values['imsi']
        key = values['key']
        opc = values['opc']
        dnn = values['dnn']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[13] = str(bandgnb)
            info[14] = str(scsgnb)
            info[15] = str(bwgnb)
            info[16] = str(mccgnb)
            info[17] = str(mncgnb)
            info[18] = str(sstgnb)
            info[19] = str(sdgnb)
            info[20] = idgnb
            info[21] = imsi
            info[22] = key
            info[23] = opc
            info[24] = dnn
        if oldinfo:
            bandgnb = info[13]
            scsgnb = info[14]
            bwgnb = info[15]
            mccgnb = info[16]
            mncgnb = info[17]
            sstgnb = info[18]
            sdgnb = info[19]
            idgnb = info[20]
            imsi = info[21]
            key = info[22]
            opc = info[23]
            dnn = info[24]
            oldinfo = False
        janela42.hide()
        if allinone and oaiue:
            janela46 = interface46()
            preparar = True
            show5 = True
            janela46['setuptype'].update(value='The setup is an All-in-One with OAI UE terminal')
        if coregnb and oaiue:
            nic_core = nic_information(ipcore, usercore, passcore)
            nic_gnb = nic_information(ipgnb, usergnb, passgnb)
            janela45 = interface45()

    if window == janela43 and event == 'Next':
        namegnb = values['namegnb']
        bandgnb = values['bandgnb']
        scsgnb = values['scsgnb']
        bwgnb = values['bwgnb']
        mccgnb = values['mccgnb']
        mncgnb = values['mncgnb']
        sstgnb = values['sstgnb']
        sdgnb = values['sdgnb']
        idgnb = values['idgnb']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[13] = str(bandgnb)
            info[14] = str(scsgnb)
            info[15] = str(bwgnb)
            info[16] = str(mccgnb)
            info[17] = str(mncgnb)
            info[18] = str(sstgnb)
            info[19] = str(sdgnb)
            info[20] = idgnb
        if oldinfo:
            bandgnb = int(info[13])
            scsgnb = int(info[14])
            bwgnb = int(info[15])
            mccgnb = int(info[16])
            mncgnb = int(info[17])
            sstgnb = int(info[18])
            sdgnb = int(info[19])
            idgnb = info[20]
            oldinfo = False
        janela43.hide()
        janela44 = interface44()

    if window == janela44 and event == 'Next':
        if allinone and realue:
            janela44.hide()
            janela46 = interface46()
            preparar = True
            show6 = True
            janela46['setuptype'].update(value='The setup is an All-in-One with a COTS UE terminal')
        if coregnb and realue:
            janela44.hide()
            nic_core = nic_information(ipcore, usercore, passcore)
            nic_gnb = nic_information(ipgnb, usergnb, passgnb)
            adm = values['adm']
            imsi = values['imsi']
            key = values['key']
            opc = values['opc']
            spn = values['spn']
            dnn = 'oai'
            program_card(adm, imsi, key, opc, spn)
            janela45 = interface45()

    if window == janela45 and event == 'Next':
        niccore = values['niccore']
        nicgnb = values['nicgnb']
        janela45.hide()
        janela46 = interface46()
        preparar = True
        if oaiue:
            print('Entra certo!!!!!!')
            threading.Thread(target=core_config_ad).start()
            threading.Thread(target=gnb_config_ad).start()
            threading.Thread(target=oaiue_config_ad).start()
            show7 = True
            janela46['setuptype'].update(value='Setup is Core and gNB on different machines with an OAI UE terminal')
            threading.Thread(target=run_core).start()
            threading.Thread(target=run_gnb).start()
            threading.Thread(target=run_oaiue).start()
        if realue:
            print('Aqui2!!!!!!!!')
            threading.Thread(target=core_config_ad).start()
            threading.Thread(target=gnb_config_ad).start()
            show8 = True
            janela46['setuptype'].update(value='Setup is Core and gNB on different machines with a COTS UE terminal')
            threading.Thread(target=run_core).start()
            threading.Thread(target=run_gnb).start()

    if window == janela47 and event == 'Finished':
        janela47.hide()
        janela1.un_hide()
        if oaiue:
            threading.Thread(target=stop_oaiue).start()
        threading.Thread(target=stop_gnb).start()
        threading.Thread(target=stop_core).start()
        show1 = False
        show2 = False
        show3 = False
        show4 = False
        show5 = False
        show6 = False
        show7 = False
        show8 = False
        limpar()

    if window == janela36 and event == 'Back':
        janela36.hide()
        janela1.un_hide()

    if window == janela37 and event == 'Back':
        janela37.hide()
        janela36.un_hide()

    if window == janela38 and event == 'Back':
        janela38.hide()
        janela37.un_hide()

    if window == janela39 and event == 'Back':
        janela39.hide()
        janela37.un_hide()

    if window == janela40 and event == 'Back':
        janela40.hide()
        janela37.un_hide()

    if window == janela41 and event == 'Back':
        janela41.hide()
        janela37.un_hide()

    if window == janela42 and event == 'Back':
        janela42.hide()
        if allinone and oaiue:
            janela38.un_hide()
        if gnbcore and oaiue:
            janela40.un_hide()

    if window == janela43 and event == 'Back':
        janela43.hide()
        if allinone and realue:
            janela39.un_hide()
        if gnbcore and realue:
            janela41.un_hide()

    if window == janela44 and event == 'Back':
        janela44.hide()
        janela43.un_hide()

    if window == janela45 and event == 'Back':
        janela45.hide()
        if gnbcore and oaiue:
            janela42.un_hide()
        if gnbcore and realue:
            janela44.un_hide()
    ################################
    #             Testes           #
    ################################

    # configuração

    if window == janela48 and event == sg.WIN_CLOSED:
        break

    if window == janela49 and event == sg.WIN_CLOSED:
        break

    if window == janela50 and event == sg.WIN_CLOSED:
        break

    if window == janela51 and event == sg.WIN_CLOSED:
        break

    if window == janela52 and event == sg.WIN_CLOSED:
        break

    if window == janela53 and event == sg.WIN_CLOSED:
        break

    if window == janela1 and event == 'Next' and values['tests']:
        tests = values['tests']
        janela1.hide()
        janela48 = interface48()

    if window == janela48 and event == 'Next' and values['config']:
        config = values['config']
        janela48.hide()
        janela49 = interface49()
        janela49['last_ipcore'].update(value=info[0])
        janela49['last_ipgnb'].update(value=info[3])
        janela49['last_ipoaiue'].update(value=info[6])

    if window == janela49 and event == 'Next':
        ipcore = values['ipcore']
        ipgnb = values['ipgnb']
        ipoaiue = values['ipue']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[0] = ipcore
            info[3] = ipgnb
            info[6] = ipoaiue
        if oldinfo:
            ipcore = info[0]
            ipgnb = info[3]
            ipoaiue = info[6]
            oldinfo = False
        janela49.hide()
        janela50 = interface50()
        janela50['last_usercore'].update(value=info[1])
        #janela50['last_passcore'].update(value=info[2])
        janela50['last_passcore'].update(value='********')
        janela50['last_usergnb'].update(value=info[4])
        #janela50['last_passgnb'].update(value=info[5])
        janela50['last_passgnb'].update(value='********')
        janela50['last_useroaiue'].update(value=info[7])
        #janela50['last_passoaiue'].update(value=info[8])
        janela50['last_passoaiue'].update(value='********')

    if window == janela50 and event == 'Next':
        usercore = values['usercore']
        passcore = values['passcore']
        usergnb = values['usergnb']
        passgnb = values['passgnb']
        useroaiue = values['userue']
        passoaiue = values['passue']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[1] = usercore
            info[2] = passcore
            info[4] = usergnb
            info[5] = passgnb
            info[7] = useroaiue
            info[8] = passoaiue
        if oldinfo:
            usercore = info[1]
            passcore = info[2]
            usergnb = info[4]
            passgnb = info[5]
            useroaiue = info[7]
            passoaiue = info[8]
            oldinfo = False
        janela50.hide()
        janela51 = interface51()
        janela51['last_dircore'].update(value=info[9])
        janela51['last_dirgnb'].update(value=info[10])
        janela51['last_diroaiue'].update(value=info[11])

    if window == janela51 and event == 'Next':
        dircore = values['dircore']
        dirgnb = values['dirgnb']
        dirue = values['dirue']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[9] = dircore
            info[10] = dirgnb
            info[11] = dirue
        if oldinfo:
            dircore = info[9]
            dirgnb = info[10]
            dirue = info[11]
            oldinfo = False
        janela51.hide()
        janela52 = interface52()
        janela52['last_gnbname'].update(value=info[12])
        janela52['last_gnbid'].update(value=info[20])
        janela52['last_mcc'].update(value=info[16])
        janela52['last_mnc'].update(value=info[17])
        janela52['last_sst'].update(value=info[18])
        janela52['last_sd'].update(value=info[19])
        janela52['last_amfip'].update(value=info[26])
        janela52['last_gnbnameamf'].update(value=info[27])
        janela52['last_gnbipamf'].update(value=info[28])
        janela52['last_gnbnamegnu'].update(value=info[29])
        janela52['last_gnbipngu'].update(value=info[30])

    if window == janela52 and event == 'Next':
        gnbname = values['gnbname']
        gnbid = values['gnbid']
        mccgnb = values['mcc']
        mncgnb = values['mnc']
        sstgnb = values['sst']
        sdgnb = values['sd']
        amfip = values['amfip']
        gnbnameamf = values['gnbnameamf']
        gnbipamf = values['gnbipamf']
        gnbnamegnu = values['gnbnamegnu']
        gnbipgnu = values['gnbipngu']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[12] = str(gnbname)
            info[25] = str(gnbid)
            info[16] = str(mccgnb)
            info[17] = str(mncgnb)
            info[18] = str(sstgnb)
            info[19] = str(sdgnb)
            info[26] = str(amfip)
            info[27] = str(gnbnameamf)
            info[28] = str(gnbipamf)
            info[29] = str(gnbnamegnu)
            info[30] = str(gnbipgnu)
        if oldinfo:
            gnbname = info[12]
            gnbid = info[25]
            mccgnb = int(info[16])
            mncgnb = int(info[17])
            sstgnb = int(info[18])
            sdgnb = int(info[19])
            amfip = info[26]
            gnbnameamf = info[27]
            gnbipamf = info[28]
            gnbnamegnu = info[29]
            gnbipgnu = info[30]
            oldinfo = False
        janela52.hide()
        janela53 = interface53()

    if window == janela53 and event == 'Save':
        janela53.hide()
        janela48.un_hide()

    if window == janela48 and event == 'Back':
        janela48.hide()
        janela1.un_hide()

    if window == janela49 and event == 'Back':
        janela49.hide()
        janela48.un_hide()

    if window == janela50 and event == 'Back':
        janela50.hide()
        janela49.un_hide()

    if window == janela51 and event == 'Back':
        janela51.hide()
        janela50.un_hide()

    if window == janela52 and event == 'Back':
        janela52.hide()
        janela51.un_hide()

    if window == janela53 and event == 'Back':
        janela53.hide()
        janela52.un_hide()

    # Testes
    if window == janela54 and event == sg.WIN_CLOSED:
        break

    if window == janela55 and event == sg.WIN_CLOSED:
        break

    if window == janela56 and event == sg.WIN_CLOSED:
        break

    if window == janela57 and event == sg.WIN_CLOSED:
        break

    if window == janela58 and event == sg.WIN_CLOSED:
        break

    if window == janela48 and event == 'Next' and values['tests']:
        tests = values['tests']
        janela48.hide()
        janela54 = interface54()

    if window == janela54 and event == 'Next' and values['coreon']:
        print('Entra')
        coreon = values['coreon']
        reltest = values['reltest']
        diretorio = values['diretorio']
        threading.Thread(target=coreon_test).start()
        janela54.hide()
        janela55 = interface55()

    if coreon:
        janela55['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)

    if preparar_ueon:
        janela55['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)

    if preparar_ueconnect:
        janela55['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)

    if preparar_iperf:
        janela55['-GIF-'].update_animation(sg.DEFAULT_BASE64_LOADING_GIF, time_between_frames=100)

    if os.path.isfile('coreon_result.txt'):

        with open('coreon_result.txt', 'r') as file:
            content = file.read()

        if ' 0% packet loss' in content:
            result_test_coreon = True
        if '100% packet loss' in content:
            result_test_coreon = False

        janela55.close()
        janela56 = interface56()

        if result_test_coreon:
            relatorio_coreconnect_certo(diretorio)
        else:
            relatorio_coreconnect_erro(diretorio)

        os.remove('coreon_result.txt')

    if result_test_coreon:
        janela56['ipcoreout'].update(value=ipcore)
        janela56['ipamfout'].update(value='192.168.70.132')
        janela56['ipranout'].update(value=ipgnb)
        result_test_coreon = False

    if window == janela56 and event == 'Close':
        janela56.hide()
        janela48.un_hide()
        coreon = False
        ueon = False

    if window == janela54 and event == 'Next' and values['ueon']:
        ueon = values['ueon']
        relgnb = values['relgnb']
        relue = values['relue']
        reltest = values['reltest']
        diretorio = values['diretorio']
        janela54.hide()
        janela57 = interface57()

    if window == janela57 and event == 'Start' and ueon:
        band = values['band']
        scs = values['scs']
        bw = values['bw']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[13] = str(band)
            info[14] = str(scs)
            info[15] = str(bw)
        if oldinfo:
            band = int(info[13])
            scs = int(info[14])
            bw = int(info[15])
            oldinfo = False
        preparar_ueon = ueon
        # Correr o gNB e o UE
        threading.Thread(target=run_gnb_tests).start()
        threading.Thread(target=run_ue_tests).start()
        threading.Thread(target=ueon_test).start()
        janela57.hide()
        janela55 = interface55()

    if os.path.isfile('ueon_result.txt'):

        with open('ueon_result.txt', 'r') as file:
            content = file.read()

        with open('ueon_result.txt', 'r') as file:
            ip = file.read()

        ips = open('ips.txt', 'r')
        lines = ips.readlines()
        for line in lines:
            if line in ip:
                result_test_ueon = True
                print(line)
                break

        stop_gnb()
        stop_oaiue()
        janela55.close()
        janela58 = interface58()
        preparar_ueon = False

        coreon = False

        # Aqui fazer o relatorio do ueon
        if reltest:
            if result_test_ueon:
                print('Relatorio do UE ON, certo')
                relatorio_ueconnect_certo(diretorio, band, scs, bw, line)
            else:
                print('Relatorio do UE ON, errado')
                relatorio_ueconnect_erro(diretorio, band, scs, bw)

        os.remove('ueon_result.txt')

    if result_test_ueon:
        janela58['ipcoreout'].update(value=ipcore)
        janela58['ipamfout'].update(value='192.168.70.132')
        janela58['ipranout'].update(value=ipgnb)
        janela58['bandout'].update(value=band)
        janela58['scsout'].update(value=scs)
        janela58['bwout'].update(value=bw)
        janela58['ipueoutmachine'].update(value=ipoaiue)
        janela58['ipueout'].update(value=line)

    if window == janela58 and event == 'Close':
        janela58.hide()
        janela48.un_hide()
        coreon = False
        ueon = False
        result_test_ueon = False

    if window == janela57 and event == 'Back':
        janela57.hide()
        janela48.un_hide()

    if window == janela54 and event == 'Back':
        janela54.hide()
        janela48.un_hide()

    if window == janela59 and event == sg.WIN_CLOSED:
        break

    if window == janela54 and event == 'Next' and values['uetoext']:
        ueconnect = values['uetoext']
        relgnb = values['relgnb']
        relue = values['relue']
        reltest = values['reltest']
        diretorio = values['diretorio']
        janela54.hide()
        janela57 = interface57()

    if window == janela57 and event == 'Start' and ueconnect:
        band = values['band']
        scs = values['scs']
        bw = values['bw']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[13] = str(band)
            info[14] = str(scs)
            info[15] = str(bw)
        if oldinfo:
            band = int(info[13])
            scs = int(info[14])
            bw = int(info[15])
            oldinfo = False
        preparar_ueconnect = ueconnect
        # Correr o gNB e o UE
        threading.Thread(target=run_gnb_tests).start()
        threading.Thread(target=run_ue_tests).start()
        threading.Thread(target=ueconnect_test).start()
        janela57.hide()
        janela55 = interface55()

    if os.path.isfile('ueconnect_result.txt'):
        print('Entrou para parar')
        time.sleep(10)
        with open('ueconnect_result.txt', 'r') as file:
            ip = file.read()

        ips = open('ips.txt', 'r')
        lines = ips.readlines()
        for line in lines:
            line = line.replace('\n', '') # fundamental
            print(line)
            if line in ip:
                print(line)
                break
        with open('ueconnect_result.txt', 'r') as file:
            data = file.read()
        if ' 0% packet loss' in data:
            result_test_ueconnect = True
        print(result_test_ueconnect)
        janela55.close()
        janela59 = interface59()

        stop_gnb()
        stop_oaiue()
        preparar_ueconnect = False
        ueconnect = False

        # Aqui fazer o relatorio do ueon
        if reltest:
            if result_test_ueconnect:
                print('Relatorio do UE Connect, certo')
                relatorio_uetoext_certo(diretorio, band, scs, bw, line)
            else:
                print('Relatorio do UE Connect, errado')
                relatorio_uetoext_erro(diretorio, band, scs, bw)

        if result_test_ueconnect:
            janela59['ipcoreout'].update(value=ipcore)
            janela59['ipamfout'].update(value='192.168.70.132')
            janela59['ipranout'].update(value=ipgnb)
            janela59['bandout'].update(value=band)
            janela59['scsout'].update(value=scs)
            janela59['bwout'].update(value=bw)
            janela59['ipueoutmachine'].update(value=ipoaiue)
            janela59['ipueout'].update(value=line)
            result_test_ueconnect = False

        os.remove('ueconnect_result.txt')

    if window == janela59 and event == 'Close':
        janela59.hide()
        janela48.un_hide()
        coreon = False
        ueon = False
        result_test_ueon = False

    if window == janela60 and event == sg.WIN_CLOSED:
        break

    if window == janela54 and event == 'Next' and values['iperftest']:
        iperf = values['iperftest']
        relgnb = values['relgnb']
        relue = values['relue']
        reltest = values['reltest']
        diretorio = values['diretorio']
        janela54.hide()
        janela57 = interface57()

    if window == janela57 and event == 'Start' and iperf:
        band = values['band']
        scs = values['scs']
        bw = values['bw']
        oldinfo = values['oldinfo']
        if not oldinfo:
            info[13] = str(band)
            info[14] = str(scs)
            info[15] = str(bw)
        if oldinfo:
            band = int(info[13])
            scs = int(info[14])
            bw = int(info[15])
            oldinfo = False
        preparar_iperf = iperf
        # Correr o gNB e o UE
        threading.Thread(target=run_gnb_tests).start()
        threading.Thread(target=run_ue_tests).start()
        threading.Thread(target=iperf_test).start()
        threading.Thread(target=iperf_ue).start()
        janela57.hide()
        janela55 = interface55()

    if os.path.isfile('stop.txt'):

        print(ipcore)

        stop_gnb()
        stop_oaiue()

        time.sleep(5) # 20
        with open('result_iperf.txt', 'r') as file:
            data = file.read()
        print(data)

        if not '0.0-20.0 sec' in data:
            result_test_iperf = False
            janela55.close()
            janela60 = interface60()
        else:
            with open('result_iperf.txt', 'r') as file:
                ip = file.read()
            ips = open('ips.txt', 'r')
            lines = ips.readlines()
            for line in lines:
                if line in ip:
                    ipue = line
                    break

            print(line)

            with open('result_iperf.txt', 'r') as file:
                data = file.read()
            if 'UDP' in data:
                tipo = 'UDP'
        #        janela4['tipotrafego'].update(value=tipo)
                print(tipo)

            if 'TCP' in data:
                tipo = 'TCP'
        #       janela4['tipotrafego'].update(value=tipo)
                print(tipo)

            lines = open('result_iperf.txt', 'r')
            count = 0
            for line in lines:
                count = count + 1
                if '0.0-20.0 sec' in line:
                    ler = line
                    result_test_iperf = True
                    break

            position_duplo = [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]
            position_barra = [0, 0]
            position_parenteses = [0, 0]
            position_parentesess = [0, 0]

            i = 0
            for p in re.finditer('  ', ler):
                position_duplo[i] = p.start()
                i = i + 1

            quantidade = ler[position_duplo[2] + 2:position_duplo[3]]
            print(quantidade)

            i = 0
            for p in re.finditer('/', ler):
                position_barra[i] = p.start()
                i = i + 1

            perdidos = ler[position_barra[1]-5 : position_barra[1]]
            print(perdidos)

            i = 0
            for p in re.finditer('\(', ler):
                position_parenteses[i] = p.start()
                i = i + 1

            total = ler[position_barra[1] + 1:position_parenteses[0]]
            print(total)

            i = 0
            for p in re.finditer('\)', ler):
                position_parentesess[i] = p.start()
                i = i + 1

            percentagem = ler[position_parenteses[0] + 1:position_parentesess[0]]
            print(percentagem)

            janela55.close()
            janela60 = interface60()

            janela60['ipcoreout'].update(value=ipcore)
            janela60['ipranout'].update(value=ipgnb)
            janela60['ipamfout'].update(value='192.168.70.132')
            janela60['bandout'].update(value=band)
            janela60['scsout'].update(value=scs)
            janela60['bwout'].update(value=bw)
            janela60['ipueoutmachine'].update(value=ipoaiue)
            janela60['ipueout'].update(value=ipue)
            janela60['tipotrafego'].update(value=tipo)
            janela60['quantidade'].update(value=quantidade)
            janela60['perdidos'].update(value=perdidos)
            janela60['total'].update(value=total)
            janela60['percentagem'].update(value=percentagem)

        if reltest:
            if result_test_iperf:
                print('Relatorio do UE Iperf, certo')
                relatorio_iperftest_certo(diretorio, band, scs, bw, line, tipo, quantidade, perdidos, total,percentagem)
            else:
                print('Relatorio do UE Iperf, errado')
                relotiro_iperftest_erro(diretorio, band, scs, bw)

        if os.path.isfile('result_iperf.txt'):
            os.remove('result_iperf.txt')
        preparar_iperf = False
        os.remove('stop.txt')

    if window == janela60 and event == 'Close':
        janela60.hide()
        janela48.un_hide()
        coreon = False
        ueon = False
        iperf = False
        result_test_ueon = False
        result_test_iperf = False
        command = 'nada'


limpar()
print(info)

if os.path.isfile('info.txt'):
    os.remove('info.txt')
with open('info.txt', 'w') as file:
    for line in info:
        file.write("".join(line) + "\n")





